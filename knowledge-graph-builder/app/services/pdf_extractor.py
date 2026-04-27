"""
PDF and DOCX extraction service.

Extracts text content (plus embedded images) from PDF and DOCX files.

Libraries used:
- PyMuPDF (fitz)   — primary PDF parser; text + image extraction
- pdfplumber       — table extraction (optional, gracefully skipped if absent)
- python-docx      — DOCX text and table extraction
- pytesseract      — OCR fallback for scanned pages (optional, gracefully skipped)
"""

from pathlib import Path
from typing import Any

from app.core.logging import get_logger

logger = get_logger(__name__)

# Minimum image dimensions to bother extracting (avoids tiny decorative images).
_MIN_IMAGE_PX = 50

# OCR trigger threshold: pages with fewer characters than this will be OCR'd.
_OCR_CHAR_THRESHOLD = 100

# Diagram detection threshold: OCR text shorter than this on an image-bearing page
# signals a likely diagram/figure rather than readable text.
_DIAGRAM_CHAR_THRESHOLD = 50

# Attempt to import OCR dependencies once at module load time.
# pytesseract is always defined at module scope (None when unavailable) so
# tests can patch it regardless of whether the package is installed.
pytesseract = None  # type: ignore[assignment]
_OCR_AVAILABLE = False

try:
    import pytesseract  # noqa: F811  (intentional re-assignment)
    from PIL import Image as _PilImage

    _OCR_AVAILABLE = True
except ImportError:
    logger.warning("pytesseract not available, OCR skipped")


def _page_to_pil(page: Any) -> Any:
    """Render a PyMuPDF page to a PIL Image for OCR."""
    pixmap = page.get_pixmap()
    from PIL import Image  # already confirmed importable when _OCR_AVAILABLE is True

    return Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)


def extract_pdf(file_path: str) -> dict[str, Any]:
    """
    Extract text and embedded images from a PDF file.

    For pages where PyMuPDF extracts fewer than 100 characters, Tesseract OCR
    is applied as a fallback (if pytesseract is installed). Pages that still
    yield very short text after OCR AND contain embedded images are flagged
    as ``likely_diagram`` for downstream structured extraction (TASK-025).

    Args:
        file_path: Absolute path to the PDF.

    Returns:
        {
            "text": str,            # full text across all pages
            "page_count": int,
            "has_tables": bool,
            "image_paths": list[str],  # PNG files written alongside the source
            "metadata": dict,
            "pages": list[dict],    # per-page metadata (ocr_used, likely_diagram, …)
        }
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError(
            "PyMuPDF is required for PDF extraction. "
            "Install it with: pip install pymupdf"
        )

    doc = fitz.open(file_path)
    page_count = doc.page_count
    text_parts: list[str] = []
    image_paths: list[str] = []
    has_tables = False
    pages_meta: list[dict] = []
    ocr_used_any = False

    img_dir = Path(file_path).parent / "extracted_images"
    img_dir.mkdir(parents=True, exist_ok=True)

    # ── Per-page text + image extraction ─────────────────────────────────────
    for page_num, page in enumerate(doc):
        page_text = page.get_text("text")
        page_images = page.get_images(full=True)

        page_meta: dict[str, Any] = {
            "page_number": page_num + 1,
            "ocr_used": False,
            "likely_diagram": False,
        }

        # ── OCR fallback ──────────────────────────────────────────────────────
        if len(page_text.strip()) < _OCR_CHAR_THRESHOLD:
            if _OCR_AVAILABLE:
                try:
                    pil_image = _page_to_pil(page)
                    ocr_text = pytesseract.image_to_string(pil_image)
                    page_text = page_text + ocr_text
                    page_meta["ocr_used"] = True
                    ocr_used_any = True
                except Exception as exc:
                    logger.warning(f"OCR failed on page {page_num + 1}: {exc}")
            # If _OCR_AVAILABLE is False, module-load warning was already emitted.

            # ── Diagram detection ─────────────────────────────────────────────
            # After OCR, if text is still very short AND the page has images,
            # mark it as a likely diagram for TASK-025 to handle.
            if len(page_text.strip()) < _DIAGRAM_CHAR_THRESHOLD and page_images:
                page_meta["likely_diagram"] = True

        if page_text.strip():
            text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")

        # ── Embedded image extraction ─────────────────────────────────────────
        for img_index, img_ref in enumerate(page_images):
            xref = img_ref[0]
            base_image = doc.extract_image(xref)
            if (
                base_image.get("width", 0) < _MIN_IMAGE_PX
                or base_image.get("height", 0) < _MIN_IMAGE_PX
            ):
                continue  # skip decorative/tiny images

            img_path = img_dir / f"page{page_num + 1}_img{img_index}.png"
            with open(img_path, "wb") as fh:
                fh.write(base_image["image"])
            image_paths.append(str(img_path))

        pages_meta.append(page_meta)

    doc.close()

    # ── Optional table extraction via pdfplumber ──────────────────────────────
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    has_tables = True
                    rows = [
                        "\t".join(str(cell) if cell else "" for cell in row)
                        for row in table
                        if row
                    ]
                    text_parts.append("[Table]\n" + "\n".join(rows))
    except ImportError:
        logger.debug("pdfplumber not available — table extraction skipped")
    except Exception as exc:
        logger.warning(f"Table extraction failed: {exc}")

    processing_method = "pymupdf+tesseract" if ocr_used_any else "pymupdf"

    return {
        "text": "\n\n".join(text_parts),
        "page_count": page_count,
        "has_tables": has_tables,
        "image_paths": image_paths,
        "pages": pages_meta,
        "metadata": {
            "page_count": page_count,
            "content_type": "application/pdf",
            "processing_method": processing_method,
            "has_embedded_images": len(image_paths) > 0,
            "ocr_used": ocr_used_any,
        },
    }


def extract_docx(file_path: str) -> dict[str, Any]:
    """
    Extract text and tables from a DOCX file using python-docx.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        {
            "text": str,
            "metadata": dict,
        }
    """
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )

    doc = docx.Document(file_path)
    text_parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    for table in doc.tables:
        rows = [
            "\t".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ]
        non_empty = [r for r in rows if r.strip()]
        if non_empty:
            text_parts.append("[Table]\n" + "\n".join(non_empty))

    return {
        "text": "\n\n".join(text_parts),
        "metadata": {
            "content_type": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            "processing_method": "python-docx",
        },
    }
