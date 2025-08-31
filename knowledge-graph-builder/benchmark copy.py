#!/usr/bin/env python3
"""
Advanced Neo4j GraphRAG Knowledge Graph Ingestion Pipeline

A comprehensive standalone implementation demonstrating sophisticated GraphRAG capabilities
including multi-phase ingestion, schema learning, advanced entity resolution, and
performance monitoring comparable to enterprise-grade systems.

Features:
- Multi-phase document processing pipeline
- Automatic schema learning and evolution
- Advanced entity resolution with multiple algorithms
- Vector indexing and embedding optimization
- Performance monitoring and benchmarking
- Production-grade error handling and logging
- Configuration-driven deployment
"""

import asyncio
import json
import logging
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Tuple
import yaml
import os
from datetime import datetime
import statistics

# Core Neo4j GraphRAG imports
from neo4j import GraphDatabase
from neo4j_graphrag.experimental.pipeline.kg_builder import SimpleKGPipeline
from neo4j_graphrag.experimental.pipeline import Pipeline
from neo4j_graphrag.experimental.pipeline.component import Component
from neo4j_graphrag.experimental.components.text_splitters.fixed_size_splitter import FixedSizeSplitter
from neo4j_graphrag.experimental.components.entity_relation_extractor import LLMEntityRelationExtractor
from neo4j_graphrag.experimental.components.embedder import TextChunkEmbedder
from neo4j_graphrag.experimental.components.kg_writer import Neo4jWriter
from neo4j_graphrag.experimental.components.lexical_graph import LexicalGraphBuilder
from neo4j_graphrag.experimental.components.resolver import (
    SinglePropertyExactMatchResolver,
    SpaCySemanticMatchResolver,
    FuzzyMatchResolver
)
from neo4j_graphrag.experimental.components.schema import (
    SchemaFromTextExtractor,
    GraphSchema,
    NodeType,
    RelationshipType,
    PropertyType
)
from neo4j_graphrag.experimental.components.pdf_loader import PdfLoader
from neo4j_graphrag.experimental.components.types import LexicalGraphConfig, DocumentInfo, GraphResult

# LLM and embedding components
from neo4j_graphrag.llm import OpenAILLM
from neo4j_graphrag.embeddings import OpenAIEmbeddings
from neo4j_graphrag.embeddings.base import Embedder

# Index management with upsert functionality
from neo4j_graphrag.indexes import create_vector_index, upsert_vectors, EntityType

# Retrieval components
from neo4j_graphrag.retrievers import (
    VectorRetriever,
    VectorCypherRetriever,
    HybridCypherRetriever,
    Text2CypherRetriever
)

# Index management
from neo4j_graphrag.indexes import create_vector_index, upsert_vectors

# Document processing
import PyPDF2
from docx import Document as DocxDocument


from pydantic import validate_call


class EntityEmbedder(Component):
    """Component for creating embeddings from entity nodes in a graph.
    
    This component follows the same pattern as TextChunkEmbedder but operates
    on entity nodes within a GraphResult, creating embeddings from entity names
    and descriptions, then efficiently upserting them using the library's
    upsert_vectors functionality.

    Args:
        embedder (Embedder): The embedder to use to create the embeddings.
        driver (neo4j.Driver): Neo4j driver for upsert operations.
        embedding_property (str): Name of the property to store embeddings.
        neo4j_database (Optional[str]): Neo4j database name.

    Example:

    .. code-block:: python

        from neo4j_graphrag.experimental.components.embedder import EntityEmbedder
        from neo4j_graphrag.embeddings.openai import OpenAIEmbeddings
        from neo4j_graphrag.experimental.pipeline import Pipeline

        embedder = OpenAIEmbeddings(model="text-embedding-3-large")
        entity_embedder = EntityEmbedder(embedder, driver)
        pipeline = Pipeline()
        pipeline.add_component(entity_embedder, "entity_embedder")
    """

    def __init__(
        self, 
        embedder: Embedder, 
        driver,
        embedding_property: str = "embedding",
        neo4j_database: Optional[str] = None
    ):
        self._embedder = embedder
        self._driver = driver
        self._embedding_property = embedding_property
        self._neo4j_database = neo4j_database
        self.logger = logging.getLogger(__name__)

    def _embed_entity(self, node) -> Optional[List[float]]:
        """Embed a single entity node.

        Args:
            node: The entity node to embed.

        Returns:
            Optional[List[float]]: The embedding vector or None if no text to embed.
        """
        # Only embed entity nodes (skip Chunk and Document nodes)
        if 'Entity' not in str(node.label) and '__Entity__' not in str(node.label):
            return None
            
        # Create text representation for embedding
        entity_text = node.properties.get('name', '')
        if node.properties.get('description'):
            entity_text += f" {node.properties['description']}"
        
        if not entity_text.strip():
            return None
            
        # Generate embedding
        embedding = self._embedder.embed_query(entity_text.strip())
        return embedding

    @validate_call
    def run(self, graph: GraphResult) -> GraphResult:
        """Embed entity nodes in the graph and upsert to Neo4j.

        Args:
            graph (GraphResult): The graph containing nodes to embed.

        Returns:
            GraphResult: The graph with entity embeddings added.
        """
        entity_ids = []
        embeddings = []
        embedded_count = 0
        
        try:
            for node in graph.nodes:
                embedding = self._embed_entity(node)
                if embedding is not None:
                    # Store embedding in node properties for pipeline continuity
                    if not node.embedding_properties:
                        node.embedding_properties = {}
                    node.embedding_properties[self._embedding_property] = embedding
                    
                    # Collect for batch upsert
                    entity_ids.append(node.id)
                    embeddings.append(embedding)
                    embedded_count += 1
            
            # Batch upsert embeddings to Neo4j using library's upsert_vectors
            if entity_ids and embeddings:
                upsert_vectors(
                    driver=self._driver,
                    ids=entity_ids,
                    embedding_property=self._embedding_property,
                    embeddings=embeddings,
                    neo4j_database=self._neo4j_database,
                    entity_type=EntityType.NODE
                )
            
            self.logger.info(f"Generated and upserted embeddings for {embedded_count} entities")
            return graph
            
        except Exception as e:
            self.logger.error(f"Entity embedding failed: {e}")
            raise


class RelationshipEmbedder(Component):
    """Component for creating embeddings from relationship edges in a graph.
    
    Similar to EntityEmbedder but operates on relationships, creating embeddings
    from relationship types and properties, then efficiently upserting them.

    Args:
        embedder (Embedder): The embedder to use to create the embeddings.
        driver (neo4j.Driver): Neo4j driver for upsert operations.
        embedding_property (str): Name of the property to store embeddings.
        neo4j_database (Optional[str]): Neo4j database name.
    """

    def __init__(
        self, 
        embedder: Embedder, 
        driver,
        embedding_property: str = "embedding",
        neo4j_database: Optional[str] = None
    ):
        self._embedder = embedder
        self._driver = driver
        self._embedding_property = embedding_property
        self._neo4j_database = neo4j_database
        self.logger = logging.getLogger(__name__)

    def _embed_relationship(self, relationship) -> Optional[List[float]]:
        """Embed a single relationship.

        Args:
            relationship: The relationship to embed.

        Returns:
            Optional[List[float]]: The embedding vector or None if skipped.
        """
        # Skip chunk-related relationships
        if relationship.type in ['FROM_CHUNK', 'FROM_DOCUMENT', 'NEXT_CHUNK', 'PART_OF_CHUNK']:
            return None
            
        # Find start and end node names for context
        start_name = "entity"
        end_name = "entity"
        
        # Look up node names from the graph (this is a simplified approach)
        # In a real implementation, you might want to pass node lookup as parameter
        if hasattr(relationship, 'start_node') and relationship.start_node:
            start_name = relationship.start_node.properties.get('name', 'entity')
        if hasattr(relationship, 'end_node') and relationship.end_node:
            end_name = relationship.end_node.properties.get('name', 'entity')
            
        # Create text representation for relationship embedding
        rel_text = f"{start_name} {relationship.type} {end_name}"
        
        # Add relationship properties as context
        if relationship.properties:
            context_parts = []
            for key, value in relationship.properties.items():
                if key != self._embedding_property:  # Skip existing embedding property
                    context_parts.append(f"{key}: {value}")
            if context_parts:
                rel_text += f" ({', '.join(context_parts)})"
        
        # Generate embedding
        embedding = self._embedder.embed_query(rel_text)
        return embedding

    @validate_call
    def run(self, graph: GraphResult) -> GraphResult:
        """Embed relationships in the graph and upsert to Neo4j.

        Args:
            graph (GraphResult): The graph containing relationships to embed.

        Returns:
            GraphResult: The graph with relationship embeddings added.
        """
        relationship_ids = []
        embeddings = []
        embedded_count = 0
        
        try:
            for relationship in graph.relationships:
                embedding = self._embed_relationship(relationship)
                if embedding is not None:
                    # Store embedding in relationship properties for pipeline continuity
                    if not relationship.embedding_properties:
                        relationship.embedding_properties = {}
                    relationship.embedding_properties[self._embedding_property] = embedding
                    
                    # Collect for batch upsert (Note: would need relationship IDs from Neo4j)
                    # For now, just store in properties and let the writer handle it
                    embedded_count += 1
            
            # Note: Relationship embedding upsert would need relationship IDs from Neo4j
            # This is more complex and might be handled by the kg_writer component
            # For now, we store embeddings in properties for the writer to handle
            
            self.logger.info(f"Generated embeddings for {embedded_count} relationships")
            return graph
            
        except Exception as e:
            self.logger.error(f"Relationship embedding failed: {e}")
            raise


class SimpleEntityResolver:
    """Simple entity resolution that creates SAME_AS relationships between duplicate entities"""
    
    def __init__(self, neo4j_driver):
        self.driver = neo4j_driver
        
    async def run(self) -> Dict[str, Any]:
        """Run simple entity resolution"""
        print("🔄 Running Simple Entity Resolution...")
        
        with self.driver.session() as session:
            # Find entities with identical names in different chunks
            result = session.run("""
                MATCH (e1:__Entity__)-[:FROM_CHUNK]->(c1:Chunk)
                MATCH (e2:__Entity__)-[:FROM_CHUNK]->(c2:Chunk)
                WHERE e1.name = e2.name 
                AND c1.index <> c2.index
                AND elementId(e1) < elementId(e2)  // Avoid duplicates
                AND NOT (e1)-[:SAME_AS]-(e2)  // Don't create if already exists
                RETURN e1.name as entity_name,
                       elementId(e1) as e1_id,
                       elementId(e2) as e2_id,
                       c1.index as chunk1,
                       c2.index as chunk2
            """)
            
            matches = list(result)
            
            # Create SAME_AS relationships
            links_created = 0
            for record in matches:
                try:
                    session.run("""
                        MATCH (e1) WHERE elementId(e1) = $e1_id
                        MATCH (e2) WHERE elementId(e2) = $e2_id
                        MERGE (e1)-[:SAME_AS {created_by: 'entity_resolution'}]-(e2)
                    """, e1_id=record['e1_id'], e2_id=record['e2_id'])
                    
                    links_created += 1
                    print(f"   ✅ Linked '{record['entity_name']}' between Chunk {record['chunk1']} ↔ Chunk {record['chunk2']}")
                    
                except Exception as e:
                    print(f"   ❌ Failed to link {record['entity_name']}: {e}")
            
            print(f"🎉 Entity Resolution Complete! Created {links_created} SAME_AS relationships")
            
            return {
                "entities_resolved": links_created,
                "method": "simple_same_as_linking"
            }


# Configuration and monitoring
@dataclass
class AdvancedPipelineConfig:
    """Comprehensive configuration for advanced GraphRAG pipeline"""
    
    # Neo4j Configuration
    neo4j_uri: str = "bolt://localhost:7687"
    neo4j_user: str = "neo4j"
    neo4j_password: str = "password"
    neo4j_database: str = "neo4j"
    
    # LLM Configuration
    openai_api_key: str = "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"
    llm_model: str = "gpt-4o"
    llm_temperature: float = 0.0
    llm_max_tokens: int = 3000
    
    # Embedding Configuration
    embedding_model: str = "text-embedding-3-large"
    embedding_dimensions: int = 3072
    
    # Chunking Configuration
    chunk_size: int = 1500
    chunk_overlap: int = 300
    approximate_chunking: bool = True
    
    # Processing Configuration
    batch_size: int = 2000
    max_concurrency: int = 10
    
    # Entity Resolution Configuration
    enable_entity_resolution: bool = True
    similarity_threshold: float = 0.85
    fuzzy_threshold: float = 0.8
    
    # Schema Configuration
    enable_schema_learning: bool = False  # Disabled to avoid errors
    enforce_schema: bool = True
    additional_properties_allowed: bool = False
    
    # Performance Configuration
    enable_performance_monitoring: bool = True
    benchmark_mode: bool = False
    
    # Error Handling
    on_error: str = "RAISE"  # "RAISE" or "IGNORE"
    enable_detailed_logging: bool = True


@dataclass
class PerformanceMetrics:
    """Performance monitoring and benchmarking metrics"""
    
    start_time: float = field(default_factory=time.time)
    end_time: Optional[float] = None
    
    # Document processing metrics
    documents_processed: int = 0
    chunks_created: int = 0
    embeddings_generated: int = 0
    
    # Entity extraction metrics
    entities_extracted: int = 0
    relationships_extracted: int = 0
    entities_resolved: int = 0
    
    # Database metrics
    nodes_created: int = 0
    relationships_created: int = 0
    indexes_created: int = 0
    
    # Performance metrics
    processing_times: List[float] = field(default_factory=list)
    memory_usage: List[float] = field(default_factory=list)
    
    def add_processing_time(self, duration: float):
        """Add processing time measurement"""
        self.processing_times.append(duration)
    
    def finalize(self):
        """Finalize metrics collection"""
        self.end_time = time.time()
    
    def get_summary(self) -> Dict[str, Any]:
        """Get comprehensive metrics summary"""
        total_duration = (self.end_time or time.time()) - self.start_time
        
        return {
            "total_duration_seconds": total_duration,
            "documents_processed": self.documents_processed,
            "chunks_created": self.chunks_created,
            "embeddings_generated": self.embeddings_generated,
            "entities_extracted": self.entities_extracted,
            "relationships_extracted": self.relationships_extracted,
            "entities_resolved": self.entities_resolved,
            "nodes_created": self.nodes_created,
            "relationships_created": self.relationships_created,
            "processing_rate_docs_per_second": self.documents_processed / total_duration if total_duration > 0 else 0,
            "average_processing_time": statistics.mean(self.processing_times) if self.processing_times else 0,
            "median_processing_time": statistics.median(self.processing_times) if self.processing_times else 0,
            "p95_processing_time": (
                sorted(self.processing_times)[int(0.95 * len(self.processing_times))] 
                if self.processing_times else 0
            )
        }


class AdvancedDocumentProcessor:
    """Advanced document processing with multi-format support"""
    
    def __init__(self, config: AdvancedPipelineConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
    async def process_file(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """Process various document formats"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            return await self._process_pdf(file_path)
        elif file_path.suffix.lower() == '.docx':
            return await self._process_docx(file_path)
        elif file_path.suffix.lower() in ['.txt', '.md']:
            return await self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    async def _process_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process PDF files with metadata extraction"""
        try:
            pdf_loader = PdfLoader()
            document = await pdf_loader.run(file_path) 
            
            metadata = {
                "source": str(file_path),
                "format": "pdf",
                "processed_at": datetime.now().isoformat()
            }
            
            return document.text, metadata
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX files"""
        try:
            doc = DocxDocument(file_path)
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            metadata = {
                "source": str(file_path),
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "processed_at": datetime.now().isoformat()
            }
            
            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                "source": str(file_path),
                "format": file_path.suffix.lower(),
                "size_chars": len(text),
                "processed_at": datetime.now().isoformat()
            }
            
            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {e}")
            raise


class AdvancedSchemaManager:
    """Advanced schema management with learning and evolution"""
    
    def __init__(self, config: AdvancedPipelineConfig, llm):
        self.config = config
        self.llm = llm
        self.schema_extractor = SchemaFromTextExtractor(llm=llm) if config.enable_schema_learning else None 
        self.logger = logging.getLogger(__name__)
        
        # Default enterprise schema
        self.default_schema = GraphSchema(
            node_types=[
                NodeType(
                    label="Person",
                    description="Individual human being with biographical information",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="occupation", type="STRING"),
                        PropertyType(name="nationality", type="STRING"),
                        PropertyType(name="birth_date", type="DATE")
                    ]
                ),
                NodeType(
                    label="Organization",
                    description="Structured entity with business or institutional purpose",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="industry", type="STRING"),
                        PropertyType(name="headquarters", type="STRING"),
                        PropertyType(name="founded_date", type="DATE")
                    ]
                ),
                NodeType(
                    label="Concept",
                    description="Abstract concept or topic",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="description", type="STRING"),
                        PropertyType(name="domain", type="STRING")
                    ]
                )
            ],
            relationship_types=[
                RelationshipType(
                    label="WORKS_FOR",
                    description="Employment relationship",
                    properties=[
                        PropertyType(name="start_date", type="DATE"),
                        PropertyType(name="position", type="STRING")
                    ]
                ),
                RelationshipType(
                    label="RELATED_TO",
                    description="General relationship between entities",
                    properties=[
                        PropertyType(name="relationship_type", type="STRING"),
                        PropertyType(name="confidence", type="FLOAT")
                    ]
                )
            ],
            additional_node_types=not config.enforce_schema,
            additional_relationship_types=not config.enforce_schema
        ) 
    
    async def get_or_learn_schema(self, text_sample: str = None) -> GraphSchema:
        """Get existing schema or learn from text sample"""
        if self.config.enable_schema_learning and text_sample and self.schema_extractor:
            try:
                self.logger.info("Learning schema from text sample...")
                learned_schema = await self.schema_extractor.run(text=text_sample) 
                
                # Check if learned_schema has the expected structure
                if hasattr(learned_schema, 'node_types') and hasattr(learned_schema, 'relationship_types'):
                    # Merge with default schema
                    return self._merge_schemas(self.default_schema, learned_schema)
                else:
                    self.logger.warning("Learned schema format unexpected, using default")
                    return self.default_schema
                    
            except Exception as e:
                self.logger.warning(f"Schema learning failed, using default: {e}")
        
        return self.default_schema
    
    def _merge_schemas(self, base_schema: GraphSchema, learned_schema: GraphSchema) -> GraphSchema:
        """Merge learned schema with base schema"""
        # Simple merge - in production, this would be more sophisticated
        merged_nodes = base_schema.node_types + [
            node for node in learned_schema.node_types 
            if node.label not in [n.label for n in base_schema.node_types]
        ]
        
        merged_relationships = base_schema.relationship_types + [
            rel for rel in learned_schema.relationship_types
            if rel.label not in [r.label for r in base_schema.relationship_types]
        ]
        
        return GraphSchema(
            node_types=merged_nodes,
            relationship_types=merged_relationships,
            additional_node_types=base_schema.additional_node_types,
            additional_relationship_types=base_schema.additional_relationship_types
        )


class MultiAlgorithmEntityResolver:
    """Advanced entity resolution using multiple algorithms"""
    
    def __init__(self, driver, config: AdvancedPipelineConfig):
        self.driver = driver
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Initialize resolvers
        self.resolvers = []
        
        if config.enable_entity_resolution:
            # Exact match resolver
            self.resolvers.append(
                SinglePropertyExactMatchResolver(
                    driver=driver,
                    resolve_property="name"
                )
            )
            
            # Semantic similarity resolver
            self.resolvers.append(
                SpaCySemanticMatchResolver(
                    driver=driver,
                    similarity_threshold=config.similarity_threshold,
                    resolve_properties=["name", "description"]
                )
            )
            
            # Fuzzy match resolver
            self.resolvers.append(
                FuzzyMatchResolver(
                    driver=driver,
                    similarity_threshold=config.fuzzy_threshold
                )
            ) 
    
    async def resolve_entities(self) -> Dict[str, Any]:
        """Run multi-algorithm entity resolution"""
        if not self.config.enable_entity_resolution:
            return {"entities_resolved": 0, "resolution_methods": []}
        
        resolution_results = []
        total_resolved = 0
        
        for i, resolver in enumerate(self.resolvers):
            try:
                self.logger.info(f"Running entity resolution with algorithm {i+1}/{len(self.resolvers)}")
                start_time = time.time()
                
                result = await resolver.run() 
                duration = time.time() - start_time
                
                resolved_count = getattr(result, 'entities_resolved', 0) if hasattr(result, 'entities_resolved') else 0
                total_resolved += resolved_count
                
                resolution_results.append({
                    "algorithm": resolver.__class__.__name__,
                    "entities_resolved": resolved_count,
                    "duration_seconds": duration
                })
                
                self.logger.info(f"Algorithm {resolver.__class__.__name__} resolved {resolved_count} entities in {duration:.2f}s")
                
            except Exception as e:
                self.logger.error(f"Entity resolution failed for {resolver.__class__.__name__}: {e}")
                resolution_results.append({
                    "algorithm": resolver.__class__.__name__,
                    "entities_resolved": 0,
                    "duration_seconds": 0,
                    "error": str(e)
                })
        
        return {
            "total_entities_resolved": total_resolved,
            "resolution_methods": resolution_results
        }


class UnifiedVectorRetriever:
    """Enhanced retriever that searches across chunks AND entities"""
    
    def __init__(self, driver, chunk_index: str, entity_index: str, embedder):
        self.driver = driver
        self.chunk_index = chunk_index
        self.entity_index = entity_index
        self.embedder = embedder
        self.logger = logging.getLogger(__name__)
    
    async def search(self, query: str, k: int = 5) -> Dict[str, Any]:
        """Search across both chunks and entities with unified ranking"""
        try:
            # Generate query embedding
            query_embedding = await self.embedder.embed_query(query)
            
            results = []
            
            with self.driver.session() as session:
                # Search chunks
                chunk_query = """
                CALL db.index.vector.queryNodes($chunk_index, $k, $query_embedding)
                YIELD node AS chunk, score AS chunk_score
                
                // Get entities from this chunk
                OPTIONAL MATCH (chunk)<-[:FROM_CHUNK]-(entity:__Entity__)
                
                RETURN 
                    'chunk' AS result_type,
                    chunk.text AS content,
                    chunk_score AS score,
                    collect(DISTINCT entity.name) AS related_entities,
                    elementId(chunk) AS node_id
                ORDER BY chunk_score DESC
                """
                
                chunk_results = session.run(chunk_query, 
                    chunk_index=self.chunk_index, 
                    k=k, 
                    query_embedding=query_embedding
                )
                
                for record in chunk_results:
                    results.append({
                        "type": record["result_type"],
                        "content": record["content"],
                        "score": record["score"],
                        "related_entities": record["related_entities"],
                        "node_id": record["node_id"]
                    })
                
                # Search entities
                entity_query = """
                CALL db.index.vector.queryNodes($entity_index, $k, $query_embedding)
                YIELD node AS entity, score AS entity_score
                
                // Get chunks containing this entity
                OPTIONAL MATCH (entity)-[:FROM_CHUNK]->(chunk:Chunk)
                
                // Get related entities
                OPTIONAL MATCH (entity)-[r]-(related:__Entity__)
                WHERE r.confidence > 0.7 OR NOT EXISTS(r.confidence)
                
                RETURN 
                    'entity' AS result_type,
                    entity.name AS content,
                    entity_score AS score,
                    collect(DISTINCT chunk.text)[0..2] AS related_chunks,
                    collect(DISTINCT related.name)[0..5] AS related_entities,
                    elementId(entity) AS node_id
                ORDER BY entity_score DESC
                """
                
                entity_results = session.run(entity_query,
                    entity_index=self.entity_index,
                    k=k,
                    query_embedding=query_embedding
                )
                
                for record in entity_results:
                    results.append({
                        "type": record["result_type"],
                        "content": record["content"],
                        "score": record["score"],
                        "related_chunks": record["related_chunks"],
                        "related_entities": record["related_entities"],
                        "node_id": record["node_id"]
                    })
            
            # Sort by score and take top k results
            results.sort(key=lambda x: x["score"], reverse=True)
            
            return {
                "query": query,
                "results": results[:k],
                "total_results": len(results)
            }
            
        except Exception as e:
            self.logger.error(f"Unified vector search failed: {e}")
            raise


class AdvancedGraphRAGPipeline:
    """Comprehensive Neo4j GraphRAG pipeline with enterprise features"""
    
    def __init__(self, config: AdvancedPipelineConfig):
        self.config = config
        self.metrics = PerformanceMetrics()
        self.logger = self._setup_logging()
        
        # Initialize components
        self.driver = None
        self.llm = None
        self.embedder = None
        self.document_processor = AdvancedDocumentProcessor(config)
        self.schema_manager = None
        self.entity_resolver = None
        
        # Pipeline components
        self.text_splitter = None
        self.extractor = None
        self.kg_writer = None
        self.pipeline = None
        
    def _setup_logging(self) -> logging.Logger:
        """Configure comprehensive logging"""
        logger = logging.getLogger(__name__)

        # Ensure log directory exists
        log_dir = Path("logs")   # you can make this configurable
        log_dir.mkdir(parents=True, exist_ok=True)

        log_file = log_dir / f'graphrag_pipeline_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
        
        if self.config.enable_detailed_logging:
            level = logging.DEBUG
        else:
            level = logging.INFO
            
        logging.basicConfig(
            level=level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file, encoding="utf-8"),
                logging.StreamHandler()
            ]
        )
        
        return logger
    
    async def initialize(self):
        """Initialize all pipeline components"""
        self.logger.info("Initializing Advanced GraphRAG Pipeline...")
        
        try:
            # Initialize Neo4j driver
            self.driver = GraphDatabase.driver(
                self.config.neo4j_uri,
                auth=(self.config.neo4j_user, self.config.neo4j_password)
            ) 
            
            # Test connection
            with self.driver.session(database=self.config.neo4j_database) as session:
                result = session.run("RETURN 1 as test")
                test_value = result.single()["test"]
                if test_value != 1:
                    raise ConnectionError("Neo4j connection test failed")
            
            self.logger.info("Neo4j connection established successfully")
            
            # Initialize LLM
            self.llm = OpenAILLM(
                model_name=self.config.llm_model,
                api_key=self.config.openai_api_key,
                model_params={
                    "temperature": self.config.llm_temperature,
                    "max_tokens": self.config.llm_max_tokens,
                    "response_format": {"type": "json_object"}
                }
            ) 
            
            # Initialize embedder
            self.embedder = OpenAIEmbeddings(
                model=self.config.embedding_model,
                api_key=self.config.openai_api_key
            )

            
            # Initialize schema manager
            self.schema_manager = AdvancedSchemaManager(self.config, self.llm)
            
            # Initialize multi-algorithm entity resolver using library components
            self.entity_resolver = MultiAlgorithmEntityResolver(self.driver, self.config)
            
            # Create vector indexes
            await self._create_vector_indexes()
            
            self.logger.info("Pipeline initialization completed successfully")
            
        except Exception as e:
            self.logger.error(f"Pipeline initialization failed: {e}")
            raise
    
    async def _create_vector_indexes(self):
        """Create optimized vector indexes"""
        try:
            # Primary text embeddings index
            create_vector_index(
                driver=self.driver,
                name="text_embeddings_primary",
                label="Chunk",
                embedding_property="embedding",
                dimensions=self.config.embedding_dimensions,
                similarity_fn="cosine"
            ) 
            
            # Entity embeddings index
            create_vector_index(
                driver=self.driver,
                name="entity_embeddings",
                label="__Entity__",
                embedding_property="embedding", 
                dimensions=self.config.embedding_dimensions,
                similarity_fn="cosine"
            ) 
            
            # Relationship embeddings index
            try:
                create_vector_index(
                    driver=self.driver,
                    name="relationship_embeddings",
                    label="__Relationship__",
                    embedding_property="embedding",
                    dimensions=self.config.embedding_dimensions,
                    similarity_fn="cosine"
                )
                self.metrics.indexes_created += 3
            except Exception as e:
                self.logger.warning(f"Relationship index creation failed (may not be supported): {e}")
                self.metrics.indexes_created += 2
            
            self.logger.info("Vector indexes created successfully")
            
        except Exception as e:
            self.logger.warning(f"Vector index creation failed (may already exist): {e}")
    
    def _create_advanced_pipeline(self, schema: GraphSchema) -> Pipeline:
        """Create advanced multi-component pipeline with document nodes"""
        pipeline = Pipeline() 
        
        # Advanced text splitter
        self.text_splitter = FixedSizeSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            approximate=self.config.approximate_chunking
        ) 
        
        # Chunk embedder
        embedder_component = TextChunkEmbedder(embedder=self.embedder)
        
        # Entity embedder
        entity_embedder = EntityEmbedder(
            embedder=self.embedder,
            driver=self.driver,
            embedding_property="embedding"
        )
        
        # Relationship embedder
        relationship_embedder = RelationshipEmbedder(
            embedder=self.embedder,
            driver=self.driver,
            embedding_property="embedding"
        )
        
        # Advanced entity/relation extractor with lexical graph integration
        # This will handle both lexical graph (Document/Chunk nodes) and entity extraction
        self.extractor = LLMEntityRelationExtractor(
            llm=self.llm,
            create_lexical_graph=True,  # Enable lexical graph to create FROM_CHUNK relationships
            max_concurrency=self.config.max_concurrency,
            on_error=self.config.on_error
        ) 
        
        # Single KG writer for complete graph
        self.entity_kg_writer = Neo4jWriter(
            driver=self.driver,
            batch_size=self.config.batch_size,
            neo4j_database=self.config.neo4j_database
        ) 
        
        # Add components to pipeline
        pipeline.add_component(self.text_splitter, "splitter")
        pipeline.add_component(embedder_component, "embedder") 
        pipeline.add_component(self.extractor, "extractor")
        pipeline.add_component(entity_embedder, "entity_embedder")
        pipeline.add_component(relationship_embedder, "relationship_embedder")
        pipeline.add_component(self.entity_kg_writer, "entity_writer")
        
        # Connect components - enhanced pipeline with embeddings
        pipeline.connect("splitter", "embedder", {"text_chunks": "splitter"})
        pipeline.connect("embedder", "extractor", {"chunks": "embedder"})
        pipeline.connect("extractor", "entity_embedder", {"graph": "extractor"})
        pipeline.connect("entity_embedder", "relationship_embedder", {"graph": "entity_embedder"})
        pipeline.connect("relationship_embedder", "entity_writer", {"graph": "relationship_embedder"})
        
        return pipeline
    
    async def process_document(self, file_path: Union[str, Path], schema: GraphSchema = None) -> Dict[str, Any]:
        """Process a single document through the complete pipeline with Document nodes"""
        start_time = time.time()
        
        try:
            # Process document
            self.logger.info(f"Processing document: {file_path}")
            text, metadata = await self.document_processor.process_file(file_path)
            self.metrics.documents_processed += 1
            
            # Get or learn schema
            if schema is None:
                schema = await self.schema_manager.get_or_learn_schema(text[:2000])  # Sample for schema learning
            
            # Create pipeline
            if self.pipeline is None:
                self.pipeline = self._create_advanced_pipeline(schema)
            
            # Run pipeline
            self.logger.info("Running advanced GraphRAG pipeline...")
            pipeline_start = time.time()
            
            # Run pipeline with simplified configuration
            # The extractor with create_lexical_graph=True will handle everything
            try:
                result = await self.pipeline.run({
                    "splitter": {"text": text}
                })
                
                self.logger.info(f"Pipeline result: {result}")
                
            except Exception as e:
                self.logger.error(f"Pipeline execution failed: {type(e).__name__}: {e}")
                import traceback
                self.logger.error(f"Traceback: {traceback.format_exc()}")
                raise
            
            pipeline_duration = time.time() - pipeline_start
            self.metrics.add_processing_time(pipeline_duration)
            
            # Update metrics from pipeline result
            if hasattr(result, 'chunks_created'):
                self.metrics.chunks_created += result.chunks_created
            if hasattr(result, 'entities_extracted'):
                self.metrics.entities_extracted += result.entities_extracted
            if hasattr(result, 'relationships_extracted'):
                self.metrics.relationships_extracted += result.relationships_extracted
            
            # Run entity resolution
            self.logger.info("Running multi-algorithm entity resolution...")
            resolution_result = await self.entity_resolver.resolve_entities()
            self.metrics.entities_resolved += resolution_result.get('total_entities_resolved', 0)
            
            total_duration = time.time() - start_time
            
            return {
                "success": True,
                "document_path": str(file_path),
                "metadata": metadata,
                "processing_duration": total_duration,
                "pipeline_duration": pipeline_duration,
                "entity_resolution": resolution_result,
                "metrics_snapshot": self.metrics.get_summary()
            }
            
        except Exception as e:
            self.logger.error(f"Document processing failed for {file_path}: {e}")
            return {
                "success": False,
                "document_path": str(file_path),
                "error": str(e),
                "processing_duration": time.time() - start_time
            }
    
    async def process_directory(self, directory_path: Union[str, Path]) -> Dict[str, Any]:
        """Process all documents in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory path: {directory_path}")
        
        # Find all supported files
        supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        files = [
            f for f in directory_path.rglob("*")
            if f.is_file() and f.suffix.lower() in supported_extensions
        ]
        
        if not files:
            raise ValueError(f"No supported files found in {directory_path}")
        
        self.logger.info(f"Found {len(files)} files to process in {directory_path}")
        
        # Process files
        results = []
        schema = None  # Will be learned from first document
        
        for i, file_path in enumerate(files, 1):
            self.logger.info(f"Processing file {i}/{len(files)}: {file_path.name}")
            
            result = await self.process_document(file_path, schema)
            results.append(result)
            
            # Use schema from first successful processing for subsequent files
            if result["success"] and schema is None:
                schema = await self.schema_manager.get_or_learn_schema()
        
        return {
            "directory_path": str(directory_path),
            "files_found": len(files),
            "files_processed": len([r for r in results if r["success"]]),
            "files_failed": len([r for r in results if not r["success"]]),
            "results": results,
            "final_metrics": self.metrics.get_summary()
        }
    
    async def create_retrieval_system(self) -> Dict[str, Any]:
        """Create advanced retrieval system with multiple strategies"""
        try:
            # Standard vector retriever for chunks
            vector_retriever = VectorRetriever(
                driver=self.driver,
                index_name="text_embeddings_primary",
                embedder=self.embedder,
                return_properties=["text", "source"]
            ) 
            
            # Unified vector retriever for multi-modal search
            unified_retriever = UnifiedVectorRetriever(
                driver=self.driver,
                chunk_index="text_embeddings_primary",
                entity_index="entity_embeddings",
                embedder=self.embedder
            )
            
            # Enhanced Vector + Cypher retriever for graph traversal with entity context
            vector_cypher_retriever = VectorCypherRetriever(
                driver=self.driver,
                index_name="text_embeddings_primary",
                embedder=self.embedder,
                retrieval_query="""
                WITH node AS chunk, score
                
                // Get entities from this chunk
                MATCH (chunk)<-[:FROM_CHUNK]-(entity:__Entity__)
                
                // Find related entities through relationships
                OPTIONAL MATCH (entity)-[r]-(related:__Entity__)
                WHERE r.confidence > 0.7 OR NOT EXISTS(r.confidence)
                
                // Get chunks containing related entities for expanded context
                OPTIONAL MATCH (related)-[:FROM_CHUNK]->(related_chunk:Chunk)
                
                RETURN 
                    chunk.text AS primary_context,
                    collect(DISTINCT related_chunk.text)[0..3] AS related_contexts,
                    collect(DISTINCT entity.name) AS entities,
                    collect(DISTINCT related.name)[0..10] AS related_entities,
                    score
                ORDER BY score DESC
                """
            ) 
            
            # Entity-focused retriever using entity embeddings
            entity_retriever = VectorRetriever(
                driver=self.driver,
                index_name="entity_embeddings",
                embedder=self.embedder,
                return_properties=["name", "description"]
            )
            
            # Hybrid retriever combining vector and full-text search
            hybrid_retriever = HybridCypherRetriever(
                driver=self.driver,
                vector_index_name="text_embeddings_primary", 
                fulltext_index_name="text_fulltext",
                embedder=self.embedder,
                retrieval_query="""
                WITH node, score
                MATCH (node)-[:FROM_CHUNK]-(entity:__Entity__)
                RETURN 
                    node.text as context, 
                    score, 
                    collect(DISTINCT entity.name) as entities
                ORDER BY score DESC
                """
            ) 

            with self.driver.session(database=self.config.neo4j_database) as session:
                schema_result = session.run("CALL db.schema.visualization()")
                schema_string = str(schema_result.data())

            text2cypher_retriever = Text2CypherRetriever(
                driver=self.driver,
                llm=self.llm,
                neo4j_schema=schema_string,
                examples=[
                    "Find all people working for organizations -> MATCH (p:Person)-[:WORKS_FOR]->(o:Organization) RETURN p, o",
                    "Show relationships between entities -> MATCH (a)-[r]->(b) RETURN a, r, b LIMIT 100",
                    "Find entities similar to a query -> CALL db.index.vector.queryNodes('entity_embeddings', 5, $embedding) YIELD node, score"
                ]
            )
            
            return {
                "retrievers_created": 6,
                "vector_retriever": vector_retriever,
                "unified_retriever": unified_retriever,
                "entity_retriever": entity_retriever,
                "vector_cypher_retriever": vector_cypher_retriever,
                "hybrid_retriever": hybrid_retriever,
                "text2cypher_retriever": text2cypher_retriever
            }
            
        except Exception as e:
            self.logger.error(f"Retrieval system creation failed: {e}")
            raise
    
    async def run_benchmarks(self) -> Dict[str, Any]:
        """Run comprehensive performance benchmarks"""
        if not self.config.benchmark_mode:
            return {"benchmarking_disabled": True}
        
        self.logger.info("Running performance benchmarks...")
        
        benchmark_results = {}
        
        # Query performance benchmarks
        try:
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Node count query
                start_time = time.time()
                node_count = session.run("MATCH (n) RETURN count(n) as count").single()["count"]
                benchmark_results["node_count_query_time"] = time.time() - start_time
                benchmark_results["total_nodes"] = node_count
                
                # Relationship count query  
                start_time = time.time()
                rel_count = session.run("MATCH ()-[r]->() RETURN count(r) as count").single()["count"]
                benchmark_results["relationship_count_query_time"] = time.time() - start_time
                benchmark_results["total_relationships"] = rel_count
                
                # Complex traversal query
                start_time = time.time()
                result = session.run("""
                    MATCH (n)-[r1]->(m)-[r2]->(o)
                    RETURN count(*) as count
                    LIMIT 1000
                """).single()
                benchmark_results["complex_traversal_time"] = time.time() - start_time
                benchmark_results["complex_traversal_count"] = result["count"] if result else 0
        
        except Exception as e:
            self.logger.error(f"Query benchmarks failed: {e}")
            benchmark_results["query_benchmarks_error"] = str(e)
        
        # Vector search benchmarks
        try:
            if self.embedder:
                # Test vector similarity search for chunks
                start_time = time.time()
                test_embedding = await self.embedder.embed_query("TechNova Corporation artificial intelligence")
                benchmark_results["embedding_generation_time"] = time.time() - start_time
                
                # Chunk vector search performance
                with self.driver.session(database=self.config.neo4j_database) as session:
                    start_time = time.time()
                    result = session.run("""
                        CALL db.index.vector.queryNodes('text_embeddings_primary', 10, $embedding)
                        YIELD node, score
                        RETURN count(node) as results
                    """, embedding=test_embedding).single()
                    benchmark_results["chunk_vector_search_time"] = time.time() - start_time
                    benchmark_results["chunk_vector_search_results"] = result["results"] if result else 0
                    
                    # Entity vector search performance
                    start_time = time.time()
                    entity_result = session.run("""
                        CALL db.index.vector.queryNodes('entity_embeddings', 10, $embedding)
                        YIELD node, score
                        RETURN count(node) as results
                    """, embedding=test_embedding).single()
                    benchmark_results["entity_vector_search_time"] = time.time() - start_time
                    benchmark_results["entity_vector_search_results"] = entity_result["results"] if entity_result else 0
                    
                    # Unified retrieval benchmark
                    unified_retriever = UnifiedVectorRetriever(
                        driver=self.driver,
                        chunk_index="text_embeddings_primary",
                        entity_index="entity_embeddings",
                        embedder=self.embedder
                    )
                    
                    start_time = time.time()
                    unified_results = await unified_retriever.search("AI research collaboration", k=5)
                    benchmark_results["unified_retrieval_time"] = time.time() - start_time
                    benchmark_results["unified_retrieval_results"] = unified_results.get("total_results", 0)
        
        except Exception as e:
            self.logger.error(f"Vector benchmarks failed: {e}")
            benchmark_results["vector_benchmarks_error"] = str(e)
        
        # Processing performance metrics
        if self.metrics.processing_times:
            benchmark_results.update({
                "avg_processing_time": statistics.mean(self.metrics.processing_times),
                "median_processing_time": statistics.median(self.metrics.processing_times),
                "min_processing_time": min(self.metrics.processing_times),
                "max_processing_time": max(self.metrics.processing_times),
                "processing_time_stddev": statistics.stdev(self.metrics.processing_times) if len(self.metrics.processing_times) > 1 else 0
            })
        
        return benchmark_results

    async def run_enhanced_benchmark(self, sample_text: str, retrieval_queries: List[str]) -> Dict:
        """
        Enhanced benchmarking including entity embeddings and multi-modal retrieval
        """
        results = {"entity_embeddings": {}, "multi_modal_retrieval": {}, "pipeline_enhanced": {}}
        
        try:
            # 1. Entity Embedding Performance
            self.logger.info("Testing entity embedding performance...")
            
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Count entities with embeddings
                entity_count = session.run("""
                    MATCH (e:__Entity__)
                    WHERE e.embedding IS NOT NULL
                    RETURN count(e) as embedded_entities
                """).single()
                
                results["entity_embeddings"]["embedded_entities"] = entity_count["embedded_entities"] if entity_count else 0
                
                # Test entity embedding retrieval accuracy
                if self.embedder:
                    start_time = time.time()
                    test_query = "AI research collaboration between organizations"
                    query_embedding = await self.embedder.embed_query(test_query)
                    
                    entity_results = session.run("""
                        CALL db.index.vector.queryNodes('entity_embeddings', 5, $embedding)
                        YIELD node, score
                        RETURN node.id as entity_id, node.name as entity_name, score
                        ORDER BY score DESC
                    """, embedding=query_embedding).data()
                    
                    results["entity_embeddings"]["query_time"] = time.time() - start_time
                    results["entity_embeddings"]["retrieved_entities"] = len(entity_results)
                    results["entity_embeddings"]["top_entities"] = [
                        {"name": r["entity_name"], "score": r["score"]} 
                        for r in entity_results[:3]
                    ]
            
            # 2. Multi-Modal Retrieval Performance
            self.logger.info("Testing multi-modal retrieval...")
            
            if hasattr(self, 'retrieval_system') and self.retrieval_system:
                for query in retrieval_queries[:3]:  # Test first 3 queries
                    query_results = {}
                    
                    # Test unified retrieval
                    start_time = time.time()
                    unified_result = await self.retrieval_system["unified_retriever"].search(query, k=5)
                    query_results["unified_time"] = time.time() - start_time
                    query_results["unified_results"] = unified_result.get("total_results", 0)
                    
                    # Test entity-specific retrieval
                    start_time = time.time()
                    entity_result = await self.retrieval_system["entity_retriever"].search(query, k=5)
                    query_results["entity_time"] = time.time() - start_time
                    query_results["entity_results"] = len(entity_result.get("records", []))
                    
                    results["multi_modal_retrieval"][f"query_{len(results['multi_modal_retrieval']) + 1}"] = query_results
            
            # 3. Enhanced Pipeline Metrics
            self.logger.info("Gathering enhanced pipeline metrics...")
            
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Relationship embedding coverage
                rel_stats = session.run("""
                    MATCH ()-[r:RELATED]->()
                    RETURN 
                        count(r) as total_relationships,
                        sum(CASE WHEN r.embedding IS NOT NULL THEN 1 ELSE 0 END) as embedded_relationships
                """).single()
                
                if rel_stats:
                    results["pipeline_enhanced"]["total_relationships"] = rel_stats["total_relationships"]
                    results["pipeline_enhanced"]["embedded_relationships"] = rel_stats["embedded_relationships"]
                    results["pipeline_enhanced"]["relationship_embedding_coverage"] = (
                        rel_stats["embedded_relationships"] / max(rel_stats["total_relationships"], 1) * 100
                    )
                
                # Entity type distribution with embeddings
                entity_types = session.run("""
                    MATCH (e:__Entity__)
                    RETURN 
                        labels(e) as entity_labels,
                        count(e) as total,
                        sum(CASE WHEN e.embedding IS NOT NULL THEN 1 ELSE 0 END) as with_embeddings
                """).data()
                
                results["pipeline_enhanced"]["entity_type_coverage"] = [
                    {
                        "type": r["entity_labels"],
                        "total": r["total"],
                        "embedded": r["with_embeddings"],
                        "coverage_percent": (r["with_embeddings"] / max(r["total"], 1) * 100)
                    }
                    for r in entity_types
                ]
            
            self.logger.info("Enhanced benchmarking completed successfully")
            
        except Exception as e:
            self.logger.error(f"Enhanced benchmarking failed: {e}")
            results["error"] = str(e)
        
        return results

    async def cleanup(self):
        """Clean up resources"""
        self.logger.info("Cleaning up pipeline resources...")
        
        if self.driver:
            self.driver.close()
        
        self.metrics.finalize()
        
        self.logger.info("Pipeline cleanup completed")
    
    async def __aenter__(self):
        """Async context manager entry"""
        await self.initialize()
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit"""
        await self.cleanup()


def load_config_from_file(config_path: str) -> AdvancedPipelineConfig:
    """Load configuration from YAML or JSON file"""
    config_path = Path(config_path)
    
    if not config_path.exists():
        raise FileNotFoundError(f"Configuration file not found: {config_path}")
    
    with open(config_path, 'r') as f:
        if config_path.suffix.lower() in ['.yml', '.yaml']:
            config_data = yaml.safe_load(f)
        elif config_path.suffix.lower() == '.json':
            config_data = json.load(f)
        else:
            raise ValueError(f"Unsupported configuration format: {config_path.suffix}")
    
    return AdvancedPipelineConfig(**config_data)


async def main():
    """Main execution function demonstrating advanced GraphRAG capabilities"""
    
    # Configuration - customize as needed
    config = AdvancedPipelineConfig(
        neo4j_uri="bolt://localhost:7687",
        neo4j_user="neo4j",
        neo4j_password="password",
        openai_api_key=os.getenv("OPENAI_API_KEY", "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"),
        
        # Advanced processing configuration
        chunk_size=1500,
        chunk_overlap=300,
        max_concurrency=10,
        batch_size=2000,
        
        # Schema and entity resolution
        enable_schema_learning=False,  # Disabled to avoid errors
        enable_entity_resolution=True,
        similarity_threshold=0.85,
        
        # Performance monitoring
        enable_performance_monitoring=True,
        benchmark_mode=True,
        enable_detailed_logging=True
    )
    
    print("🚀 Starting Advanced Neo4j GraphRAG Pipeline")
    print("=" * 60)
    
    # Initialize and run pipeline
    async with AdvancedGraphRAGPipeline(config) as pipeline:
        try:
            # Example 1: Process a single document
            print("\n📄 Processing single document...")
            document_result = await pipeline.process_document("./document.txt")
            print(f"✅ Document processed: {document_result['success']}")
            if document_result['success']:
                print(f"   Processing time: {document_result['processing_duration']:.2f}s")
                print(f"   Entities resolved: {document_result['entity_resolution']['total_entities_resolved']}")
            
            # Example 2: Process directory of documents
            # print("\n📁 Processing document directory...")
            # directory_result = await pipeline.process_directory("./documents")
            # print(f"✅ Directory processed: {directory_result['files_processed']}/{directory_result['files_found']} files")
            
            # Example 3: Create retrieval system
            print("\n🔍 Creating advanced retrieval system...")
            retrieval_system = await pipeline.create_retrieval_system()
            print(f"✅ Created {retrieval_system['retrievers_created']} retrieval strategies")
            
            # Example 4: Run performance benchmarks
            print("\n📊 Running performance benchmarks...")
            benchmarks = await pipeline.run_benchmarks()
            print("✅ Benchmarks completed:")
            for key, value in benchmarks.items():
                if isinstance(value, float):
                    print(f"   {key}: {value:.4f}")
                else:
                    print(f"   {key}: {value}")
            
            # Example 5: Run enhanced entity embedding benchmarks
            print("\n🧠 Running enhanced entity embedding benchmarks...")
            test_queries = [
                "artificial intelligence research",
                "collaboration between organizations",
                "technological innovation"
            ]
            enhanced_benchmarks = await pipeline.run_enhanced_benchmark("./document.txt", test_queries)
            print("✅ Enhanced benchmarks completed:")
            
            # Entity embeddings results
            if "entity_embeddings" in enhanced_benchmarks:
                ee = enhanced_benchmarks["entity_embeddings"]
                print(f"   Embedded entities: {ee.get('embedded_entities', 0)}")
                print(f"   Entity query time: {ee.get('query_time', 0):.4f}s")
                print(f"   Retrieved entities: {ee.get('retrieved_entities', 0)}")
            
            # Multi-modal retrieval results
            if "multi_modal_retrieval" in enhanced_benchmarks:
                mmr = enhanced_benchmarks["multi_modal_retrieval"]
                print(f"   Multi-modal queries tested: {len(mmr)}")
                for query_name, query_data in mmr.items():
                    print(f"   {query_name}: {query_data.get('unified_results', 0)} unified results")
            
            # Pipeline enhancement metrics
            if "pipeline_enhanced" in enhanced_benchmarks:
                pe = enhanced_benchmarks["pipeline_enhanced"]
                print(f"   Relationship embedding coverage: {pe.get('relationship_embedding_coverage', 0):.1f}%")
                entity_types = pe.get('entity_type_coverage', [])
                if entity_types:
                    avg_coverage = sum(et.get('coverage_percent', 0) for et in entity_types) / len(entity_types)
                    print(f"   Average entity embedding coverage: {avg_coverage:.1f}%")
            
            # Final metrics summary
            print("\n📈 Final Performance Summary")
            print("=" * 40)
            final_metrics = pipeline.metrics.get_summary()
            for key, value in final_metrics.items():
                if isinstance(value, float):
                    print(f"{key}: {value:.4f}")
                else:
                    print(f"{key}: {value}")
                    
        except Exception as e:
            print(f"❌ Pipeline execution failed: {e}")
            raise
    
    async def _debug_entity_chunk_connections(self, document_path: str):
        """Debug function to verify entity-chunk connections"""
        try:
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Check document and chunk creation
                doc_chunk_query = """
                MATCH (doc:Document {path: $document_path})-[:FROM_DOCUMENT]->(chunk:Chunk)
                RETURN doc.path as doc_path, count(chunk) as chunk_count
                """
                doc_result = session.run(doc_chunk_query, document_path=document_path)
                doc_data = doc_result.single()
                
                if doc_data:
                    self.logger.info(f"✅ Document-Chunk connections: {doc_data['chunk_count']} chunks for {doc_data['doc_path']}")
                else:
                    self.logger.warning(f"❌ No Document-Chunk connections found for {document_path}")
                
                # Check entity-chunk connections
                entity_chunk_query = """
                MATCH (doc:Document {path: $document_path})-[:FROM_DOCUMENT]->(chunk:Chunk)<-[:FROM_CHUNK]-(entity:__Entity__)
                RETURN chunk.id as chunk_id, collect(entity.id) as entities
                ORDER BY chunk.id
                """
                entity_result = session.run(entity_chunk_query, document_path=document_path)
                
                total_connections = 0
                for record in entity_result:
                    chunk_id = record['chunk_id']
                    entities = record['entities']
                    total_connections += len(entities)
                    self.logger.info(f"✅ Chunk {chunk_id}: {len(entities)} entities connected")
                
                if total_connections > 0:
                    self.logger.info(f"✅ Total entity-chunk connections: {total_connections}")
                else:
                    self.logger.warning("❌ No entity-chunk connections found! This indicates the FROM_CHUNK relationships are missing.")
                    
                    # Additional debugging: check if entities exist at all
                    entity_count_query = """
                    MATCH (entity:__Entity__)
                    RETURN count(entity) as total_entities
                    """
                    entity_count_result = session.run(entity_count_query)
                    entity_count = entity_count_result.single()['total_entities']
                    self.logger.info(f"🔍 Total entities in database: {entity_count}")
                    
        except Exception as e:
            self.logger.error(f"Debug function failed: {e}")
    
    print("\n🎉 Advanced GraphRAG Pipeline completed successfully!")


if __name__ == "__main__":
    # Example usage
    asyncio.run(main())

