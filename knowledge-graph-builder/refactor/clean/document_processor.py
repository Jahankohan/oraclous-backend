"""Document processing utilities - exact same functionality as original."""

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from neo4j_graphrag.experimental.components.pdf_loader import PdfLoader


class AdvancedDocumentProcessor:
    """Advanced document processing with multi-format support - exact copy from original"""

    def __init__(self, config):
        self.config = config
        import logging

        self.logger = logging.getLogger(__name__)

    async def process_file(self, file_path: str | Path) -> tuple[str, dict[str, Any]]:
        """Process various document formats"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract text based on file type
        if file_path.suffix.lower() == ".pdf":
            return await self._process_pdf(file_path)
        elif file_path.suffix.lower() == ".docx":
            return await self._process_docx(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            return await self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    async def _process_pdf(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process PDF files with metadata extraction"""
        try:
            pdf_loader = PdfLoader()
            document = await pdf_loader.run(file_path)

            metadata = {
                "source": str(file_path),
                "format": "pdf",
                "processed_at": datetime.now().isoformat(),
            }

            return document.text, metadata
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    async def _process_docx(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process DOCX files"""
        try:
            doc = DocxDocument(file_path)
            text = "\n\n".join(
                [
                    paragraph.text
                    for paragraph in doc.paragraphs
                    if paragraph.text.strip()
                ]
            )

            metadata = {
                "source": str(file_path),
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "processed_at": datetime.now().isoformat(),
            }

            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    async def _process_text(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process text files"""
        try:
            with open(file_path, encoding="utf-8") as f:
                text = f.read()

            metadata = {
                "source": str(file_path),
                "format": file_path.suffix.lower(),
                "size_chars": len(text),
                "processed_at": datetime.now().isoformat(),
            }

            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {e}")
            raise
