#!/usr/bin/env python3
"""
Advanced Neo4j GraphRAG Knowledge Graph Ingestion Pipeline

A comprehensive standalone implementation demonstrating sophisticated GraphRAG capabilities
including multi-phase ingestion, schema learning, advanced entity resolution, and
performance monitoring comparable to enterprise-grade systems.

Features:
- Multi-phase document processing pipeline
- Automatic schema learning and evolution
- Advanced entity resolution with multiple algorithms
- Vector indexing and embedding optimization
- Performance monitoring and benchmarking
- Production-grade error handling and logging
- Configuration-driven deployment
"""

import asyncio
import json
import logging
import os
import statistics
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

# Document processing
import yaml
from docx import Document as DocxDocument

# Core Neo4j GraphRAG imports
from neo4j import GraphDatabase
from neo4j_graphrag.embeddings import OpenAIEmbeddings
from neo4j_graphrag.experimental.components.embedder import TextChunkEmbedder
from neo4j_graphrag.experimental.components.entity_relation_extractor import (
    LLMEntityRelationExtractor,
)
from neo4j_graphrag.experimental.components.kg_writer import Neo4jWriter
from neo4j_graphrag.experimental.components.pdf_loader import PdfLoader
from neo4j_graphrag.experimental.components.resolver import (
    FuzzyMatchResolver,
    SinglePropertyExactMatchResolver,
    SpaCySemanticMatchResolver,
)
from neo4j_graphrag.experimental.components.schema import (
    GraphSchema,
    NodeType,
    PropertyType,
    RelationshipType,
    SchemaFromTextExtractor,
)
from neo4j_graphrag.experimental.components.text_splitters.fixed_size_splitter import (
    FixedSizeSplitter,
)
from neo4j_graphrag.experimental.components.types import (
    DocumentInfo,
    LexicalGraphConfig,
    Neo4jGraph,
)
from neo4j_graphrag.experimental.pipeline import Pipeline
from neo4j_graphrag.experimental.pipeline.component import Component

# Index management
from neo4j_graphrag.indexes import (
    create_fulltext_index,
    create_vector_index,
)

# LLM and embedding components
from neo4j_graphrag.llm import OpenAILLM

# Retrieval components
from neo4j_graphrag.retrievers import (
    HybridRetriever,
    Text2CypherRetriever,
    VectorCypherRetriever,
    VectorRetriever,
)
from pydantic import validate_call


class LoggingText2CypherRetriever:
    """Wrapper for Text2CypherRetriever that logs generated Cypher queries"""

    def __init__(self, text2cypher_retriever, logger):
        self._retriever = text2cypher_retriever
        self.logger = logger
        self.last_generated_cypher = None

    def search(self, query_text: str, **kwargs):
        """Override search to capture and log generated Cypher"""
        try:
            # Call the original search method
            result = self._retriever.search(query_text=query_text, **kwargs)

            # Try to capture the generated Cypher from various possible locations
            generated_cypher = None

            # Method 1: Check result metadata (most common location)
            if (
                hasattr(result, "metadata")
                and result.metadata
                and "cypher" in result.metadata
            ):
                generated_cypher = result.metadata["cypher"]
                self.logger.info("🔍 Found Cypher in result.metadata['cypher']")

            # Method 2: Check result attributes
            elif hasattr(result, "cypher_query"):
                generated_cypher = result.cypher_query
                self.logger.info("🔍 Found Cypher in result.cypher_query")
            elif hasattr(result, "query"):
                generated_cypher = result.query
                self.logger.info("🔍 Found Cypher in result.query")

            # Method 3: Check retriever attributes
            elif hasattr(self._retriever, "last_query"):
                generated_cypher = self._retriever.last_query
                self.logger.info("🔍 Found Cypher in _retriever.last_query")
            elif hasattr(self._retriever, "_last_cypher"):
                generated_cypher = self._retriever._last_cypher
                self.logger.info("🔍 Found Cypher in _retriever._last_cypher")

            # Store for access
            self.last_generated_cypher = generated_cypher

            # Log the query
            if generated_cypher:
                self.logger.info(
                    f"🔍 Generated Cypher for '{query_text}': {generated_cypher}"
                )
            else:
                self.logger.warning(
                    f"⚠️  Could not capture generated Cypher for query: '{query_text}'"
                )
                # Debug info
                if hasattr(result, "metadata"):
                    self.logger.debug(
                        f"Result metadata keys: {list(result.metadata.keys()) if result.metadata else 'None'}"
                    )

            # Enhance result with generated cypher if possible
            if hasattr(result, "__dict__") and generated_cypher:
                result.generated_cypher = generated_cypher

            return result

        except Exception as e:
            self.logger.error(
                f"❌ Error in Text2CypherRetriever for '{query_text}': {e}"
            )
            raise

    def __getattr__(self, name):
        """Delegate other attributes to the wrapped retriever"""
        return getattr(self._retriever, name)


class EntityEmbedder(Component):
    """Simple entity embedder following neo4j-graphrag patterns"""

    def __init__(self, embedder):
        """Initialize with just an embedder like TextChunkEmbedder"""
        self.embedder = embedder

    @validate_call
    async def run(self, graph: Neo4jGraph) -> Neo4jGraph:
        """Add entity embeddings to entities without existing embeddings."""
        if not graph.nodes:
            return graph

        # Find entities without embeddings (excluding Document and Chunk nodes)
        entities_to_embed = [
            node
            for node in graph.nodes
            if (
                node.label not in ["Document", "Chunk"]
                and (
                    not node.embedding_properties
                    or not node.embedding_properties.get("embedding")
                )
            )
        ]

        if entities_to_embed:
            # Create embedding text for each entity
            embedding_texts = []
            for entity in entities_to_embed:
                # Create descriptive text for entity embedding
                name = entity.properties.get("name", "")
                label = entity.label
                text = f"{label}: {name}" if name else label
                embedding_texts.append(text)

            # Get embeddings in batch
            embeddings = []
            for text in embedding_texts:
                embedding = self.embedder.embed_query(text)
                embeddings.append(embedding)

            # Add embeddings to entities
            for entity, embedding in zip(entities_to_embed, embeddings, strict=False):
                if not entity.embedding_properties:
                    entity.embedding_properties = {}
                entity.embedding_properties["embedding"] = embedding

            print(f"✅ Added embeddings to {len(entities_to_embed)} entities")

        return graph


class RelationshipEmbedder(Component):
    """Simple relationship embedder following neo4j-graphrag patterns"""

    def __init__(self, embedder):
        """Initialize with just an embedder like TextChunkEmbedder"""
        self.embedder = embedder

    @validate_call
    async def run(self, graph: Neo4jGraph) -> Neo4jGraph:
        """Add relationship embeddings to relationships without existing embeddings."""
        if not graph.relationships:
            return graph

        # Find relationships without embeddings
        rels_to_embed = [
            rel
            for rel in graph.relationships
            if (
                not rel.embedding_properties
                or not rel.embedding_properties.get("embedding")
            )
        ]

        if rels_to_embed:
            # Create embedding text for each relationship
            embedding_texts = []
            for rel in rels_to_embed:
                # Create descriptive text for relationship embedding
                start_name = rel.start_node_id
                end_name = rel.end_node_id
                rel_type = rel.type
                text = f"{start_name} {rel_type} {end_name}"
                embedding_texts.append(text)

            # Get embeddings in batch
            embeddings = []
            for text in embedding_texts:
                embedding = self.embedder.embed_query(text)
                embeddings.append(embedding)

            # Add embeddings to relationships
            for rel, embedding in zip(rels_to_embed, embeddings, strict=False):
                if not rel.embedding_properties:
                    rel.embedding_properties = {}
                rel.embedding_properties["embedding"] = embedding

            print(f"✅ Added embeddings to {len(rels_to_embed)} relationships")

        return graph


class SimpleEntityResolver:
    """Simple entity resolution that creates SAME_AS relationships between duplicate entities"""

    def __init__(self, neo4j_driver):
        self.driver = neo4j_driver

    async def run(self) -> dict[str, Any]:
        """Run simple entity resolution"""
        print("🔄 Running Simple Entity Resolution...")

        with self.driver.session() as session:
            # Find entities with identical names in different chunks
            result = session.run("""
                MATCH (e1:__Entity__)-[:FROM_CHUNK]->(c1:Chunk)
                MATCH (e2:__Entity__)-[:FROM_CHUNK]->(c2:Chunk)
                WHERE e1.name = e2.name
                AND c1.index <> c2.index
                AND elementId(e1) < elementId(e2)  // Avoid duplicates
                AND NOT (e1)-[:SAME_AS]-(e2)  // Don't create if already exists
                RETURN e1.name as entity_name,
                       elementId(e1) as e1_id,
                       elementId(e2) as e2_id,
                       c1.index as chunk1,
                       c2.index as chunk2
            """)

            matches = list(result)

            # Create SAME_AS relationships
            links_created = 0
            for record in matches:
                try:
                    session.run(
                        """
                        MATCH (e1) WHERE elementId(e1) = $e1_id
                        MATCH (e2) WHERE elementId(e2) = $e2_id
                        MERGE (e1)-[:SAME_AS {created_by: 'entity_resolution'}]-(e2)
                    """,
                        e1_id=record["e1_id"],
                        e2_id=record["e2_id"],
                    )

                    links_created += 1
                    print(
                        f"   ✅ Linked '{record['entity_name']}' between Chunk {record['chunk1']} ↔ Chunk {record['chunk2']}"
                    )

                except Exception as e:
                    print(f"   ❌ Failed to link {record['entity_name']}: {e}")

            print(
                f"🎉 Entity Resolution Complete! Created {links_created} SAME_AS relationships"
            )

            return {
                "entities_resolved": links_created,
                "method": "simple_same_as_linking",
            }


# Configuration and monitoring
@dataclass
class AdvancedPipelineConfig:
    """Comprehensive configuration for advanced GraphRAG pipeline"""

    # Neo4j Configuration
    neo4j_uri: str = "bolt://localhost:7687"
    neo4j_user: str = "neo4j"
    neo4j_password: str = ""
    neo4j_database: str = "neo4j"

    # LLM Configuration
    openai_api_key: str = ""
    llm_model: str = "gpt-4o"
    llm_temperature: float = 0.0
    llm_max_tokens: int = 3000

    # Embedding Configuration
    embedding_model: str = "text-embedding-3-large"
    embedding_dimensions: int = 3072

    # Chunking Configuration
    chunk_size: int = 1500
    chunk_overlap: int = 300
    approximate_chunking: bool = True

    # Processing Configuration
    batch_size: int = 2000
    max_concurrency: int = 10

    # Entity Resolution Configuration
    enable_entity_resolution: bool = True
    similarity_threshold: float = 0.85
    fuzzy_threshold: float = 0.8

    # Schema Configuration
    enable_schema_learning: bool = True
    enforce_schema: bool = True
    additional_properties_allowed: bool = False

    # Performance Configuration
    enable_performance_monitoring: bool = True
    benchmark_mode: bool = False

    # Error Handling
    on_error: str = "RAISE"  # "RAISE" or "IGNORE"
    enable_detailed_logging: bool = True


@dataclass
class PerformanceMetrics:
    """Performance monitoring and benchmarking metrics"""

    start_time: float = field(default_factory=time.time)
    end_time: float | None = None

    # Document processing metrics
    documents_processed: int = 0
    chunks_created: int = 0
    embeddings_generated: int = 0

    # Entity extraction metrics
    entities_extracted: int = 0
    relationships_extracted: int = 0
    entities_resolved: int = 0

    # Database metrics
    nodes_created: int = 0
    relationships_created: int = 0
    indexes_created: int = 0

    # Performance metrics
    processing_times: list[float] = field(default_factory=list)
    memory_usage: list[float] = field(default_factory=list)

    def add_processing_time(self, duration: float):
        """Add processing time measurement"""
        self.processing_times.append(duration)

    def finalize(self):
        """Finalize metrics collection"""
        self.end_time = time.time()

    def get_summary(self) -> dict[str, Any]:
        """Get comprehensive metrics summary"""
        total_duration = (self.end_time or time.time()) - self.start_time

        return {
            "total_duration_seconds": total_duration,
            "documents_processed": self.documents_processed,
            "chunks_created": self.chunks_created,
            "embeddings_generated": self.embeddings_generated,
            "entities_extracted": self.entities_extracted,
            "relationships_extracted": self.relationships_extracted,
            "entities_resolved": self.entities_resolved,
            "nodes_created": self.nodes_created,
            "relationships_created": self.relationships_created,
            "processing_rate_docs_per_second": self.documents_processed / total_duration
            if total_duration > 0
            else 0,
            "average_processing_time": statistics.mean(self.processing_times)
            if self.processing_times
            else 0,
            "median_processing_time": statistics.median(self.processing_times)
            if self.processing_times
            else 0,
            "p95_processing_time": (
                sorted(self.processing_times)[int(0.95 * len(self.processing_times))]
                if self.processing_times
                else 0
            ),
        }


class AdvancedDocumentProcessor:
    """Advanced document processing with multi-format support"""

    def __init__(self, config: AdvancedPipelineConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)

    async def process_file(self, file_path: str | Path) -> tuple[str, dict[str, Any]]:
        """Process various document formats"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract text based on file type
        if file_path.suffix.lower() == ".pdf":
            return await self._process_pdf(file_path)
        elif file_path.suffix.lower() == ".docx":
            return await self._process_docx(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            return await self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    async def _process_pdf(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process PDF files with metadata extraction"""
        try:
            pdf_loader = PdfLoader()
            document = await pdf_loader.run(file_path)

            metadata = {
                "source": str(file_path),
                "format": "pdf",
                "processed_at": datetime.now().isoformat(),
            }

            return document.text, metadata
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    async def _process_docx(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process DOCX files"""
        try:
            doc = DocxDocument(file_path)
            text = "\n\n".join(
                [
                    paragraph.text
                    for paragraph in doc.paragraphs
                    if paragraph.text.strip()
                ]
            )

            metadata = {
                "source": str(file_path),
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "processed_at": datetime.now().isoformat(),
            }

            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    async def _process_text(self, file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process text files"""
        try:
            with open(file_path, encoding="utf-8") as f:
                text = f.read()

            metadata = {
                "source": str(file_path),
                "format": file_path.suffix.lower(),
                "size_chars": len(text),
                "processed_at": datetime.now().isoformat(),
            }

            return text, metadata
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {e}")
            raise


class AdvancedSchemaManager:
    """Advanced schema management with learning and evolution"""

    def __init__(self, config: AdvancedPipelineConfig, llm):
        self.config = config
        self.llm = llm
        self.schema_extractor = (
            SchemaFromTextExtractor(llm=llm) if config.enable_schema_learning else None
        )
        self.logger = logging.getLogger(__name__)

        # Default enterprise schema
        self.default_schema = GraphSchema(
            node_types=[
                NodeType(
                    label="Person",
                    description="Individual human being with biographical information",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="occupation", type="STRING"),
                        PropertyType(name="nationality", type="STRING"),
                        PropertyType(name="birth_date", type="DATE"),
                    ],
                ),
                NodeType(
                    label="Organization",
                    description="Structured entity with business or institutional purpose",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="industry", type="STRING"),
                        PropertyType(name="headquarters", type="STRING"),
                        PropertyType(name="founded_date", type="DATE"),
                    ],
                ),
                NodeType(
                    label="Concept",
                    description="Abstract concept or topic",
                    properties=[
                        PropertyType(name="name", type="STRING", required=True),
                        PropertyType(name="description", type="STRING"),
                        PropertyType(name="domain", type="STRING"),
                    ],
                ),
            ],
            relationship_types=[
                RelationshipType(
                    label="WORKS_FOR",
                    description="Employment relationship",
                    properties=[
                        PropertyType(name="start_date", type="DATE"),
                        PropertyType(name="position", type="STRING"),
                    ],
                ),
                RelationshipType(
                    label="RELATED_TO",
                    description="General relationship between entities",
                    properties=[
                        PropertyType(name="relationship_type", type="STRING"),
                        PropertyType(name="confidence", type="FLOAT"),
                    ],
                ),
            ],
            additional_node_types=not config.enforce_schema,
            additional_relationship_types=not config.enforce_schema,
        )

    async def get_or_learn_schema(self, text_sample: str = None) -> GraphSchema:
        """Get existing schema or learn from text sample"""
        if self.config.enable_schema_learning and text_sample and self.schema_extractor:
            try:
                self.logger.info("Learning schema from text sample...")
                learned_schema = await self.schema_extractor.run(text=text_sample)

                # Merge with default schema
                return self._merge_schemas(self.default_schema, learned_schema)
            except Exception as e:
                self.logger.warning(f"Schema learning failed, using default: {e}")

        return self.default_schema

    def _merge_schemas(
        self, base_schema: GraphSchema, learned_schema: GraphSchema
    ) -> GraphSchema:
        """Merge learned schema with base schema"""
        # Simple merge - in production, this would be more sophisticated
        merged_nodes = base_schema.node_types + [
            node
            for node in learned_schema.node_types
            if node.label not in [n.label for n in base_schema.node_types]
        ]

        merged_relationships = base_schema.relationship_types + [
            rel
            for rel in learned_schema.relationship_types
            if rel.label not in [r.label for r in base_schema.relationship_types]
        ]

        return GraphSchema(
            node_types=merged_nodes,
            relationship_types=merged_relationships,
            additional_node_types=base_schema.additional_node_types,
            additional_relationship_types=base_schema.additional_relationship_types,
        )


class MultiAlgorithmEntityResolver:
    """Advanced entity resolution using multiple algorithms"""

    def __init__(self, driver, config: AdvancedPipelineConfig):
        self.driver = driver
        self.config = config
        self.logger = logging.getLogger(__name__)

        # Initialize resolvers
        self.resolvers = []

        if config.enable_entity_resolution:
            # Exact match resolver
            self.resolvers.append(
                SinglePropertyExactMatchResolver(driver=driver, resolve_property="name")
            )

            # Semantic similarity resolver
            self.resolvers.append(
                SpaCySemanticMatchResolver(
                    driver=driver,
                    similarity_threshold=config.similarity_threshold,
                    resolve_properties=["name", "description"],
                )
            )

            # Fuzzy match resolver
            self.resolvers.append(
                FuzzyMatchResolver(
                    driver=driver, similarity_threshold=config.fuzzy_threshold
                )
            )

    async def resolve_entities(self) -> dict[str, Any]:
        """Run multi-algorithm entity resolution"""
        if not self.config.enable_entity_resolution:
            return {"entities_resolved": 0, "resolution_methods": []}

        resolution_results = []
        total_resolved = 0

        for i, resolver in enumerate(self.resolvers):
            try:
                self.logger.info(
                    f"Running entity resolution with algorithm {i + 1}/{len(self.resolvers)}"
                )
                start_time = time.time()

                result = await resolver.run()
                duration = time.time() - start_time

                resolved_count = (
                    getattr(result, "entities_resolved", 0)
                    if hasattr(result, "entities_resolved")
                    else 0
                )
                total_resolved += resolved_count

                resolution_results.append(
                    {
                        "algorithm": resolver.__class__.__name__,
                        "entities_resolved": resolved_count,
                        "duration_seconds": duration,
                    }
                )

                self.logger.info(
                    f"Algorithm {resolver.__class__.__name__} resolved {resolved_count} entities in {duration:.2f}s"
                )

            except Exception as e:
                self.logger.error(
                    f"Entity resolution failed for {resolver.__class__.__name__}: {e}"
                )
                resolution_results.append(
                    {
                        "algorithm": resolver.__class__.__name__,
                        "entities_resolved": 0,
                        "duration_seconds": 0,
                        "error": str(e),
                    }
                )

        return {
            "total_entities_resolved": total_resolved,
            "resolution_methods": resolution_results,
        }


class AdvancedGraphRAGPipeline:
    """Comprehensive Neo4j GraphRAG pipeline with enterprise features"""

    def __init__(self, config: AdvancedPipelineConfig):
        self.config = config
        self.metrics = PerformanceMetrics()
        self.logger = self._setup_logging()

        # Initialize components
        self.driver = None
        self.llm = None
        self.embedder = None
        self.document_processor = AdvancedDocumentProcessor(config)
        self.schema_manager = None
        self.entity_resolver = None

        # Pipeline components
        self.text_splitter = None
        self.extractor = None
        self.kg_writer = None
        self.pipeline = None

    def _setup_logging(self) -> logging.Logger:
        """Configure comprehensive logging"""
        logger = logging.getLogger(__name__)

        # Ensure log directory exists
        log_dir = Path("logs")  # you can make this configurable
        log_dir.mkdir(parents=True, exist_ok=True)

        log_file = (
            log_dir
            / f"graphrag_pipeline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        )

        if self.config.enable_detailed_logging:
            level = logging.DEBUG
        else:
            level = logging.INFO

        logging.basicConfig(
            level=level,
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(log_file, encoding="utf-8"),
                logging.StreamHandler(),
            ],
        )

        return logger

    async def initialize(self):
        """Initialize all pipeline components"""
        self.logger.info("Initializing Advanced GraphRAG Pipeline...")

        try:
            # Initialize Neo4j driver
            self.driver = GraphDatabase.driver(
                self.config.neo4j_uri,
                auth=(self.config.neo4j_user, self.config.neo4j_password),
            )

            # Test connection
            with self.driver.session(database=self.config.neo4j_database) as session:
                result = session.run("RETURN 1 as test")
                test_value = result.single()["test"]
                if test_value != 1:
                    raise ConnectionError("Neo4j connection test failed")

            self.logger.info("Neo4j connection established successfully")

            # Initialize LLM
            self.llm = OpenAILLM(
                model_name=self.config.llm_model,
                api_key=self.config.openai_api_key,
                model_params={
                    "temperature": self.config.llm_temperature,
                    "max_tokens": self.config.llm_max_tokens,
                    "response_format": {"type": "json_object"},
                },
            )

            # Initialize embedder
            self.embedder = OpenAIEmbeddings(
                model=self.config.embedding_model, api_key=self.config.openai_api_key
            )

            # Initialize schema manager
            self.schema_manager = AdvancedSchemaManager(self.config, self.llm)

            # Initialize multi-algorithm entity resolver using library components
            self.entity_resolver = MultiAlgorithmEntityResolver(
                self.driver, self.config
            )

            # Create vector indexes
            await self._create_vector_indexes()

            self.logger.info("Pipeline initialization completed successfully")

        except Exception as e:
            self.logger.error(f"Pipeline initialization failed: {e}")
            raise

    async def _create_vector_indexes(self):
        """Create optimized vector indexes for comprehensive GraphRAG chatbot support"""
        try:
            # Primary text embeddings index for chunks
            create_vector_index(
                driver=self.driver,
                name="text_embeddings_primary",
                label="Chunk",
                embedding_property="embedding",
                dimensions=self.config.embedding_dimensions,
                similarity_fn="cosine",
            )

            # Entity embeddings index for semantic entity search
            create_vector_index(
                driver=self.driver,
                name="entity_embeddings",
                label="__Entity__",
                embedding_property="embedding",
                dimensions=self.config.embedding_dimensions,
                similarity_fn="cosine",
            )

            # Relationship embeddings index for semantic relationship search
            create_vector_index(
                driver=self.driver,
                name="relationship_embeddings",
                label="__Relationship__",
                embedding_property="embedding",
                dimensions=self.config.embedding_dimensions,
                similarity_fn="cosine",
            )

            # Schema-aware entity type indexes for better retrieval
            for entity_type in ["Person", "Organization", "Concept"]:
                try:
                    create_vector_index(
                        driver=self.driver,
                        name=f"{entity_type.lower()}_embeddings",
                        label=entity_type,
                        embedding_property="embedding",
                        dimensions=self.config.embedding_dimensions,
                        similarity_fn="cosine",
                    )
                except Exception as e:
                    self.logger.debug(
                        f"Schema-specific index for {entity_type} may not be needed: {e}"
                    )

            # Create fulltext indexes for hybrid search capabilities
            self._create_fulltext_indexes()

            self.metrics.indexes_created += 3  # Base indexes
            self.logger.info("Vector indexes created successfully")

        except Exception as e:
            self.logger.warning(
                f"Vector index creation failed (may already exist): {e}"
            )

    def _create_fulltext_indexes(self):
        """Create fulltext indexes for hybrid GraphRAG retrieval"""

        try:
            # Fulltext index for chunk text content
            create_fulltext_index(
                driver=self.driver,
                name="chunk_text_fulltext",
                label="Chunk",
                node_properties=["text"],
                neo4j_database=self.config.neo4j_database,
            )

            # Fulltext index for entity names and descriptions
            create_fulltext_index(
                driver=self.driver,
                name="entity_text_fulltext",
                label="__Entity__",
                node_properties=["name", "description"],
                neo4j_database=self.config.neo4j_database,
            )

            self.logger.info("Fulltext indexes created for hybrid search")

        except Exception as e:
            self.logger.warning(
                f"Fulltext index creation failed (may already exist): {e}"
            )

    def _create_advanced_pipeline(self, schema: GraphSchema) -> Pipeline:
        """Create advanced multi-component pipeline with semantic embeddings"""
        pipeline = Pipeline()

        # Advanced text splitter
        self.text_splitter = FixedSizeSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            approximate=self.config.approximate_chunking,
        )

        # Chunk embedder
        chunk_embedder = TextChunkEmbedder(embedder=self.embedder)

        # Advanced entity/relation extractor with lexical graph integration
        # This will handle both lexical graph (Document/Chunk nodes) and entity extraction
        self.extractor = LLMEntityRelationExtractor(
            llm=self.llm,
            create_lexical_graph=True,  # Enable lexical graph to create FROM_CHUNK relationships
            max_concurrency=self.config.max_concurrency,
            on_error=self.config.on_error,
        )

        # Semantic embedders following library patterns
        entity_embedder = EntityEmbedder(embedder=self.embedder)
        relationship_embedder = RelationshipEmbedder(embedder=self.embedder)

        # Single KG writer for complete graph
        self.entity_kg_writer = Neo4jWriter(
            driver=self.driver,
            batch_size=self.config.batch_size,
            neo4j_database=self.config.neo4j_database,
        )

        # Add components to pipeline
        pipeline.add_component(self.text_splitter, "splitter")
        pipeline.add_component(chunk_embedder, "chunk_embedder")
        pipeline.add_component(self.extractor, "extractor")
        pipeline.add_component(entity_embedder, "entity_embedder")
        pipeline.add_component(relationship_embedder, "relationship_embedder")
        pipeline.add_component(self.entity_kg_writer, "entity_writer")

        # Connect components following clean data flow
        pipeline.connect(
            "splitter", "chunk_embedder", input_config={"text_chunks": "splitter"}
        )
        pipeline.connect(
            "chunk_embedder", "extractor", input_config={"chunks": "chunk_embedder"}
        )
        pipeline.connect(
            "extractor", "entity_embedder", input_config={"graph": "extractor"}
        )
        pipeline.connect(
            "entity_embedder",
            "relationship_embedder",
            input_config={"graph": "entity_embedder"},
        )
        pipeline.connect(
            "relationship_embedder",
            "entity_writer",
            input_config={"graph": "relationship_embedder"},
        )

        return pipeline

    async def process_document(
        self, file_path: str | Path, schema: GraphSchema = None
    ) -> dict[str, Any]:
        """Process a single document through the complete pipeline with Document nodes"""
        start_time = time.time()

        try:
            # Process document
            self.logger.info(f"Processing document: {file_path}")
            text, metadata = await self.document_processor.process_file(file_path)
            self.metrics.documents_processed += 1

            # Get or learn schema
            if schema is None:
                schema = await self.schema_manager.get_or_learn_schema(
                    text[:2000]
                )  # Sample for schema learning

            # Create pipeline
            if self.pipeline is None:
                self.pipeline = self._create_advanced_pipeline(schema)

            # Create document info for entity extraction
            document_info = DocumentInfo(path=str(file_path))

            # Run pipeline
            self.logger.info("Running advanced GraphRAG pipeline...")
            pipeline_start = time.time()

            # Run pipeline with semantic embedding configuration
            result = await self.pipeline.run(
                {
                    "splitter": {"text": text},
                    "extractor": {
                        "document_info": document_info,  # Pass document info to extractor
                        "lexical_graph_config": LexicalGraphConfig(
                            id_increment_size=50, batch_size=self.config.batch_size
                        ),
                    },
                }
            )

            pipeline_duration = time.time() - pipeline_start
            self.metrics.add_processing_time(pipeline_duration)

            # Update metrics from pipeline result
            if hasattr(result, "chunks_created"):
                self.metrics.chunks_created += result.chunks_created
            if hasattr(result, "entities_extracted"):
                self.metrics.entities_extracted += result.entities_extracted
            if hasattr(result, "relationships_extracted"):
                self.metrics.relationships_extracted += result.relationships_extracted

            # Run entity resolution
            self.logger.info("Running multi-algorithm entity resolution...")
            resolution_result = await self.entity_resolver.resolve_entities()
            self.metrics.entities_resolved += resolution_result.get(
                "total_entities_resolved", 0
            )

            total_duration = time.time() - start_time

            return {
                "success": True,
                "document_path": str(file_path),
                "metadata": metadata,
                "processing_duration": total_duration,
                "pipeline_duration": pipeline_duration,
                "entity_resolution": resolution_result,
                "metrics_snapshot": self.metrics.get_summary(),
            }

        except Exception as e:
            self.logger.error(f"Document processing failed for {file_path}: {e}")
            return {
                "success": False,
                "document_path": str(file_path),
                "error": str(e),
                "processing_duration": time.time() - start_time,
            }

    async def process_directory(self, directory_path: str | Path) -> dict[str, Any]:
        """Process all documents in a directory"""
        directory_path = Path(directory_path)

        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory path: {directory_path}")

        # Find all supported files
        supported_extensions = {".pdf", ".docx", ".txt", ".md"}
        files = [
            f
            for f in directory_path.rglob("*")
            if f.is_file() and f.suffix.lower() in supported_extensions
        ]

        if not files:
            raise ValueError(f"No supported files found in {directory_path}")

        self.logger.info(f"Found {len(files)} files to process in {directory_path}")

        # Process files
        results = []
        schema = None  # Will be learned from first document

        for i, file_path in enumerate(files, 1):
            self.logger.info(f"Processing file {i}/{len(files)}: {file_path.name}")

            result = await self.process_document(file_path, schema)
            results.append(result)

            # Use schema from first successful processing for subsequent files
            if result["success"] and schema is None:
                schema = await self.schema_manager.get_or_learn_schema()

        return {
            "directory_path": str(directory_path),
            "files_found": len(files),
            "files_processed": len([r for r in results if r["success"]]),
            "files_failed": len([r for r in results if not r["success"]]),
            "results": results,
            "final_metrics": self.metrics.get_summary(),
        }

    async def create_chatbot_ready_assessment(self) -> dict[str, Any]:
        """Comprehensive assessment of GraphRAG chatbot readiness"""
        assessment = {
            "overall_readiness": "READY_WITH_ENHANCEMENTS",
            "readiness_score": 0.75,  # 75% ready
            "strengths": [],
            "missing_critical": [],
            "recommended_enhancements": [],
            "performance_metrics": {},
        }

        try:
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Check vector indexes
                vector_indexes = session.run("""
                    SHOW INDEXES YIELD name, type, labelsOrTypes, properties
                    WHERE type = 'VECTOR'
                    RETURN name, labelsOrTypes, properties
                """).data()

                # Check fulltext indexes
                fulltext_indexes = session.run("""
                    SHOW INDEXES YIELD name, type, labelsOrTypes
                    WHERE type = 'FULLTEXT'
                    RETURN name, labelsOrTypes
                """).data()

                # Check graph structure
                structure_stats = session.run("""
                    MATCH (d:Document)-[:FROM_DOCUMENT]->(c:Chunk)<-[:FROM_CHUNK]-(e:__Entity__)
                    RETURN
                        count(DISTINCT d) as documents,
                        count(DISTINCT c) as chunks,
                        count(DISTINCT e) as entities,
                        count(DISTINCT labels(e)) as entity_types
                """).single()

                # Check embeddings coverage
                embedding_coverage = session.run("""
                    MATCH (c:Chunk)
                    WITH count(c) as total_chunks
                    MATCH (c:Chunk) WHERE c.embedding IS NOT NULL
                    RETURN total_chunks, count(c) as chunks_with_embeddings,
                           round(100.0 * count(c) / total_chunks) as coverage_percentage
                """).single()

                # Assess strengths
                if len(vector_indexes) >= 2:
                    assessment["strengths"].append(
                        "Multiple vector indexes for entities and chunks"
                    )
                if structure_stats["documents"] > 0:
                    assessment["strengths"].append(
                        "Proper Document-Chunk-Entity structure"
                    )
                if embedding_coverage["coverage_percentage"] > 90:
                    assessment["strengths"].append("High embedding coverage")

                # Assess missing critical components
                if len(fulltext_indexes) == 0:
                    assessment["missing_critical"].append(
                        "Fulltext indexes for hybrid search"
                    )
                    assessment["readiness_score"] -= 0.15

                if structure_stats["entity_types"] < 3:
                    assessment["missing_critical"].append(
                        "Diverse entity types for rich retrieval"
                    )
                    assessment["readiness_score"] -= 0.10

                # Performance metrics
                assessment["performance_metrics"] = {
                    "total_documents": structure_stats["documents"],
                    "total_chunks": structure_stats["chunks"],
                    "total_entities": structure_stats["entities"],
                    "entity_types": structure_stats["entity_types"],
                    "embedding_coverage": f"{embedding_coverage['coverage_percentage']}%",
                    "vector_indexes": len(vector_indexes),
                    "fulltext_indexes": len(fulltext_indexes),
                }

                # Recommendations
                assessment["recommended_enhancements"] = [
                    "Add fulltext indexes for hybrid search capabilities",
                    "Implement entity relationship confidence scoring",
                    "Add document metadata for better context filtering",
                    "Create specialized retrievers for different query types",
                    "Implement query routing based on user intent",
                ]

                # Update overall readiness
                if assessment["readiness_score"] >= 0.9:
                    assessment["overall_readiness"] = "PRODUCTION_READY"
                elif assessment["readiness_score"] >= 0.7:
                    assessment["overall_readiness"] = "READY_WITH_MINOR_ENHANCEMENTS"
                elif assessment["readiness_score"] >= 0.5:
                    assessment["overall_readiness"] = "NEEDS_SIGNIFICANT_WORK"
                else:
                    assessment["overall_readiness"] = "NOT_READY"

        except Exception as e:
            assessment["error"] = str(e)
            assessment["overall_readiness"] = "ASSESSMENT_FAILED"

        return assessment

    async def create_retrieval_system(self) -> dict[str, Any]:
        """Create advanced retrieval system optimized for GraphRAG chatbot with multiple strategies"""
        try:
            # 1. Primary Vector Retriever for semantic chunk search
            vector_retriever = VectorRetriever(
                driver=self.driver,
                index_name="text_embeddings_primary",
                embedder=self.embedder,
                return_properties=["text", "chunk_index"],
            )

            # 2. Enhanced Vector + Cypher Retriever for graph-aware context
            vector_cypher_retriever = VectorCypherRetriever(
                driver=self.driver,
                index_name="text_embeddings_primary",
                embedder=self.embedder,
                retrieval_query="""
                WITH node AS chunk, score
                MATCH (chunk)<-[:FROM_CHUNK]-(entity:__Entity__)-[r]->(related_entity:__Entity__)
                WHERE r.confidence > 0.5
                RETURN
                    chunk.text AS context,
                    chunk.chunk_index AS chunk_index,
                    collect(DISTINCT {
                        entity: entity.name,
                        type: labels(entity)[0],
                        relationship: type(r),
                        related_entity: related_entity.name,
                        confidence: r.confidence
                    }) AS knowledge_graph_context,
                    score
                ORDER BY score DESC
                """,
            )

            # 3. Entity-focused Vector Retriever for entity-centric queries
            entity_retriever = VectorRetriever(
                driver=self.driver,
                index_name="entity_embeddings",
                embedder=self.embedder,
                return_properties=["name", "description", "type"],
            )

            # 4. Hybrid Retriever combining vector and full-text search
            hybrid_retriever = HybridRetriever(
                driver=self.driver,
                vector_index_name="text_embeddings_primary",
                fulltext_index_name="chunk_text_fulltext",
                embedder=self.embedder,
            )

            # 5. Enhanced Text2Cypher with comprehensive schema
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Get comprehensive schema information
                schema_result = session.run("""
                    CALL db.schema.visualization() YIELD nodes, relationships
                    RETURN nodes, relationships
                """)
                schema_data = schema_result.single()

                # Enhanced schema with examples
                enhanced_schema = self._build_enhanced_schema(schema_data)

            custom_prompt = """
            Generate a Cypher query for a Neo4j graph database based on this natural language request.

            Available Schema:
            {schema}

            Query Examples:
            {examples}

            User Question:
            {query_text}

            Generate a single Cypher query that answers the user's question. The query must:
            - Use only the nodes, relationships, and properties shown in the schema
            - Start with MATCH, CREATE, MERGE, or other valid Cypher keywords
            - Be syntactically correct and executable
            - Return relevant data to answer the question

            Output ONLY the raw Cypher query without any JSON formatting, code blocks, or explanations. Example:
            MATCH (n:Entity) RETURN n LIMIT 10
            """
            text2cypher_retriever = Text2CypherRetriever(
                driver=self.driver,
                llm=self.llm,
                neo4j_schema=enhanced_schema,
                custom_prompt=custom_prompt,
                examples=[
                    "Find all people working for TechNova Corporation -> MATCH (e1:__Entity__)-[r:WORKS_FOR|EMPLOYED_BY|CEO_OF|CTO_OF]->(e2:__Entity__ {name: 'TechNova'}) RETURN e1.name, type(r), e2.name",
                    "Show companies founded in Austin -> MATCH (e1:__Entity__)-[r:LOCATED_IN|FOUNDED_IN]->(e2:__Entity__ {name: 'Austin'}) RETURN e1.name, type(r)",
                    "Find entities related to artificial intelligence -> MATCH (e:__Entity__) WHERE e.name CONTAINS 'AI' OR e.name CONTAINS 'artificial intelligence' OR e.description CONTAINS 'artificial intelligence' RETURN e.name, e.description LIMIT 10",
                    "Get document context for entity -> MATCH (e:__Entity__ {name: $entity_name})-[:FROM_CHUNK]->(c:Chunk)-[:FROM_DOCUMENT]->(d:Document) RETURN d.path, c.text",
                    "Show entity relationships with high confidence -> MATCH (e1:__Entity__)-[r]->(e2:__Entity__) WHERE r.confidence > 0.8 RETURN e1.name, type(r), e2.name, r.confidence ORDER BY r.confidence DESC LIMIT 20",
                ],
            )

            # Wrap with logging to capture generated Cypher queries
            text2cypher_retriever = LoggingText2CypherRetriever(
                text2cypher_retriever, self.logger
            )

            return {
                "retrievers_created": 5,
                "vector_retriever": vector_retriever,
                "vector_cypher_retriever": vector_cypher_retriever,
                "entity_retriever": entity_retriever,
                "hybrid_retriever": hybrid_retriever,
                "text2cypher_retriever": text2cypher_retriever,
                "chatbot_ready": True,
                "supported_queries": [
                    "semantic_similarity",
                    "graph_traversal",
                    "entity_search",
                    "hybrid_search",
                    "natural_language_to_cypher",
                ],
            }

        except Exception as e:
            self.logger.error(f"Retrieval system creation failed: {e}")
            raise

    def _build_enhanced_schema(self, schema_data) -> str:
        """Build enhanced schema description for Text2Cypher"""
        schema_description = """
        # Knowledge Graph Schema for GraphRAG

        ## Core Node Types:
        - Document: Represents source documents
          Properties: path (string), metadata (map)

        - Chunk: Text chunks from documents
          Properties: text (string), chunk_index (integer), embedding (vector)

        - __Entity__: Extracted entities (people, organizations, concepts, etc.)
          Properties: name (string), description (string), type (string), embedding (vector)

        ## Key Relationships:
        - FROM_DOCUMENT: Links chunks to their source documents (Document)-[:FROM_DOCUMENT]->(Chunk)
        - FROM_CHUNK: Links entities to the chunks they were extracted from (Chunk)<-[:FROM_CHUNK]-(__Entity__)
        - Dynamic relationships between entities: (__Entity__)-[various_types]->(__Entity__)
          Common types: WORKS_FOR, LOCATED_IN, FOUNDED_BY, DEVELOPS, PARTNERS_WITH, etc.
        - SAME_AS: Entity resolution links for duplicate entities

        ## Important Notes:
        - All extracted entities use the __Entity__ label regardless of type (person, organization, concept)
        - Entity relationships have confidence scores (r.confidence property)
        - Chunks do not have a 'source' property - use path from related Document instead
        - Use vector similarity searches for semantic matching

        ## Vector Indexes Available:
        - text_embeddings_primary: For semantic search on chunks
        - entity_embeddings: For semantic search on entities

        ## Fulltext Indexes Available:
        - chunk_text_fulltext: For keyword search on chunk text
        - entity_text_fulltext: For keyword search on entity names/descriptions
        """

        return schema_description

    async def run_benchmarks(self) -> dict[str, Any]:
        """Run comprehensive performance benchmarks"""
        if not self.config.benchmark_mode:
            return {"benchmarking_disabled": True}

        self.logger.info("Running performance benchmarks...")

        benchmark_results = {}

        # Query performance benchmarks
        try:
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Node count query
                start_time = time.time()
                node_count = session.run("MATCH (n) RETURN count(n) as count").single()[
                    "count"
                ]
                benchmark_results["node_count_query_time"] = time.time() - start_time
                benchmark_results["total_nodes"] = node_count

                # Relationship count query
                start_time = time.time()
                rel_count = session.run(
                    "MATCH ()-[r]->() RETURN count(r) as count"
                ).single()["count"]
                benchmark_results["relationship_count_query_time"] = (
                    time.time() - start_time
                )
                benchmark_results["total_relationships"] = rel_count

                # Complex traversal query
                start_time = time.time()
                result = session.run("""
                    MATCH (n)-[r1]->(m)-[r2]->(o)
                    RETURN count(*) as count
                    LIMIT 1000
                """).single()
                benchmark_results["complex_traversal_time"] = time.time() - start_time
                benchmark_results["complex_traversal_count"] = (
                    result["count"] if result else 0
                )

        except Exception as e:
            self.logger.error(f"Query benchmarks failed: {e}")
            benchmark_results["query_benchmarks_error"] = str(e)

        # Vector search benchmarks
        try:
            if self.embedder:
                # Test vector similarity search
                start_time = time.time()
                test_embedding = self.embedder.embed_query(
                    "TechNova Corporation artificial intelligence"
                )
                benchmark_results["embedding_generation_time"] = (
                    time.time() - start_time
                )

                # Vector search performance
                with self.driver.session(
                    database=self.config.neo4j_database
                ) as session:
                    start_time = time.time()
                    result = session.run(
                        """
                        CALL db.index.vector.queryNodes('text_embeddings_primary', 10, $embedding)
                        YIELD node, score
                        RETURN count(node) as results
                    """,
                        embedding=test_embedding,
                    ).single()
                    benchmark_results["vector_search_time"] = time.time() - start_time
                    benchmark_results["vector_search_results"] = (
                        result["results"] if result else 0
                    )

        except Exception as e:
            self.logger.error(f"Vector benchmarks failed: {e}")
            benchmark_results["vector_benchmarks_error"] = str(e)

        # Processing performance metrics
        if self.metrics.processing_times:
            benchmark_results.update(
                {
                    "avg_processing_time": statistics.mean(
                        self.metrics.processing_times
                    ),
                    "median_processing_time": statistics.median(
                        self.metrics.processing_times
                    ),
                    "min_processing_time": min(self.metrics.processing_times),
                    "max_processing_time": max(self.metrics.processing_times),
                    "processing_time_stddev": statistics.stdev(
                        self.metrics.processing_times
                    )
                    if len(self.metrics.processing_times) > 1
                    else 0,
                }
            )

        return benchmark_results

    async def cleanup(self):
        """Clean up resources"""
        self.logger.info("Cleaning up pipeline resources...")

        if self.driver:
            self.driver.close()

        self.metrics.finalize()

        self.logger.info("Pipeline cleanup completed")

    async def __aenter__(self):
        """Async context manager entry"""
        await self.initialize()
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit"""
        await self.cleanup()


def load_config_from_file(config_path: str) -> AdvancedPipelineConfig:
    """Load configuration from YAML or JSON file"""
    config_path = Path(config_path)

    if not config_path.exists():
        raise FileNotFoundError(f"Configuration file not found: {config_path}")

    with open(config_path) as f:
        if config_path.suffix.lower() in [".yml", ".yaml"]:
            config_data = yaml.safe_load(f)
        elif config_path.suffix.lower() == ".json":
            config_data = json.load(f)
        else:
            raise ValueError(f"Unsupported configuration format: {config_path.suffix}")

    return AdvancedPipelineConfig(**config_data)


async def main():
    """Main execution function demonstrating advanced GraphRAG capabilities"""

    # Configuration - customize as needed
    config = AdvancedPipelineConfig(
        neo4j_uri="bolt://localhost:7687",
        neo4j_user="neo4j",
        neo4j_password="",
        openai_api_key=os.environ["OPENAI_API_KEY"],
        # Advanced processing configuration
        chunk_size=1500,
        chunk_overlap=300,
        max_concurrency=10,
        batch_size=2000,
        # Schema and entity resolution
        enable_schema_learning=True,
        enable_entity_resolution=True,
        similarity_threshold=0.85,
        # Performance monitoring
        enable_performance_monitoring=True,
        benchmark_mode=True,
        enable_detailed_logging=True,
    )

    print("🚀 Starting Advanced Neo4j GraphRAG Pipeline")
    print("=" * 60)

    # Initialize and run pipeline
    async with AdvancedGraphRAGPipeline(config) as pipeline:
        try:
            # Example 1: Process a single document
            print("\n📄 Processing single document...")
            document_result = await pipeline.process_document("./document.txt")
            print(f"✅ Document processed: {document_result['success']}")
            if document_result["success"]:
                print(
                    f"   Processing time: {document_result['processing_duration']:.2f}s"
                )
                print(
                    f"   Entities resolved: {document_result['entity_resolution']['total_entities_resolved']}"
                )

            # Example 2: Process directory of documents
            # print("\n📁 Processing document directory...")
            # directory_result = await pipeline.process_directory("./documents")
            # print(f"✅ Directory processed: {directory_result['files_processed']}/{directory_result['files_found']} files")

            # Example 3: Create retrieval system
            print("\n🔍 Creating advanced retrieval system...")
            retrieval_system = await pipeline.create_retrieval_system()
            print(
                f"✅ Created {retrieval_system['retrievers_created']} retrieval strategies"
            )

            # Example 4: Run performance benchmarks
            print("\n📊 Running performance benchmarks...")
            benchmarks = await pipeline.run_benchmarks()
            print("✅ Benchmarks completed:")
            for key, value in benchmarks.items():
                if isinstance(value, float):
                    print(f"   {key}: {value:.4f}")
                else:
                    print(f"   {key}: {value}")

            # Final metrics summary
            print("\n📈 Final Performance Summary")
            print("=" * 40)
            final_metrics = pipeline.metrics.get_summary()
            for key, value in final_metrics.items():
                if isinstance(value, float):
                    print(f"{key}: {value:.4f}")
                else:
                    print(f"{key}: {value}")

        except Exception as e:
            print(f"❌ Pipeline execution failed: {e}")
            raise

    async def _debug_entity_chunk_connections(self, document_path: str):
        """Debug function to verify entity-chunk connections"""
        try:
            with self.driver.session(database=self.config.neo4j_database) as session:
                # Check document and chunk creation
                doc_chunk_query = """
                MATCH (doc:Document {path: $document_path})-[:FROM_DOCUMENT]->(chunk:Chunk)
                RETURN doc.path as doc_path, count(chunk) as chunk_count
                """
                doc_result = session.run(doc_chunk_query, document_path=document_path)
                doc_data = doc_result.single()

                if doc_data:
                    self.logger.info(
                        f"✅ Document-Chunk connections: {doc_data['chunk_count']} chunks for {doc_data['doc_path']}"
                    )
                else:
                    self.logger.warning(
                        f"❌ No Document-Chunk connections found for {document_path}"
                    )

                # Check entity-chunk connections
                entity_chunk_query = """
                MATCH (doc:Document {path: $document_path})-[:FROM_DOCUMENT]->(chunk:Chunk)<-[:FROM_CHUNK]-(entity:__Entity__)
                RETURN chunk.id as chunk_id, collect(entity.id) as entities
                ORDER BY chunk.id
                """
                entity_result = session.run(
                    entity_chunk_query, document_path=document_path
                )

                total_connections = 0
                for record in entity_result:
                    chunk_id = record["chunk_id"]
                    entities = record["entities"]
                    total_connections += len(entities)
                    self.logger.info(
                        f"✅ Chunk {chunk_id}: {len(entities)} entities connected"
                    )

                if total_connections > 0:
                    self.logger.info(
                        f"✅ Total entity-chunk connections: {total_connections}"
                    )
                else:
                    self.logger.warning(
                        "❌ No entity-chunk connections found! This indicates the FROM_CHUNK relationships are missing."
                    )

                    # Additional debugging: check if entities exist at all
                    entity_count_query = """
                    MATCH (entity:__Entity__)
                    RETURN count(entity) as total_entities
                    """
                    entity_count_result = session.run(entity_count_query)
                    entity_count = entity_count_result.single()["total_entities"]
                    self.logger.info(f"🔍 Total entities in database: {entity_count}")

        except Exception as e:
            self.logger.error(f"Debug function failed: {e}")

    print("\n🎉 Advanced GraphRAG Pipeline completed successfully!")


if __name__ == "__main__":
    # Example usage
    asyncio.run(main())
